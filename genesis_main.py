# ==============================================================================
# ИМПОРТЫ
# ==============================================================================
# 1.1. Стандартная библиотека Python
import functools
import hashlib
import json
import logging
import math
import os
import queue
import random
import re
import shutil
import sys
import threading
import time
import traceback
import uuid
from dataclasses import asdict, dataclass, field
from datetime import datetime
from enum import Enum
from typing import (Any, Callable, Dict, List, Literal, Optional, Type, TypeVar,
                      Union)
# 1.2. Сторонние библиотеки (из requirements.txt)
import arxiv
import customtkinter as ctk
import docx
import networkx as nx
import numpy as np
import pandas as pd
import PyPDF2
import requests
import urllib3
import wikipedia
from bs4 import BeautifulSoup
from ddgs import DDGS
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL 
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from openai import OpenAI
from pydantic import BaseModel, Field, ValidationError, field_validator
import subprocess

# 1.3. Локальные модули проекта
from config import CONFIG, SystemConfig

# 1.4. Инициализация и настройка после импорта
# Отключаем предупреждения SSL (для локальной работы)
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# ==============================================================================
# ГЛОБАЛЬНАЯ НАСТРОЙКА ЛОГИРОВАНИЯ
# ==============================================================================

for path in [CONFIG.LOG_DIR, CONFIG.THOUGHT_DIR, CONFIG.WORK_DIR, os.path.join(CONFIG.WORK_DIR, "states")]:
    os.makedirs(path, exist_ok=True)

# 1. Основной лог системы (Технический)
LOG_FILE = os.path.join(CONFIG.LOG_DIR, f"genesis_core_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log")

log_formatter = logging.Formatter(
    '%(asctime)s | %(levelname)-8s | %(module)-15s | %(funcName)-20s | %(message)s',
    datefmt='%H:%M:%S'
)

file_handler = logging.FileHandler(LOG_FILE, encoding='utf-8')
file_handler.setLevel(logging.DEBUG)
file_handler.setFormatter(log_formatter)

console_handler = logging.StreamHandler(sys.stdout)
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(log_formatter)

logger = logging.getLogger("Genesis")
logger.setLevel(logging.DEBUG)
# Очищаем старые хендлеры, чтобы не дублировать, если код перезапускается
if logger.hasHandlers():
    logger.handlers.clear()
logger.addHandler(file_handler)
logger.addHandler(console_handler)

# 2. Логгер "Полного Потока Сознания" (FULL RAW LOG)
THOUGHT_FILE = os.path.join(CONFIG.THOUGHT_DIR, f"FULL_DEBUG_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt")

thought_logger = logging.getLogger("GenesisDeepThought")
thought_logger.setLevel(logging.INFO)
if thought_logger.hasHandlers():
    thought_logger.handlers.clear()

# Форматтер без даты внутри сообщения (она будет в заголовке блока)
thought_handler = logging.FileHandler(THOUGHT_FILE, encoding='utf-8')
thought_handler.setFormatter(logging.Formatter('%(message)s'))
thought_logger.addHandler(thought_handler)

logger.info("=== ЗАПУСК GENESIS ===")


# ==============================================================================
# 1. КОНФИГУРАЦИЯ И КОНСТАНТЫ 
# ==============================================================================

@dataclass
class MissionState:
    topic: str
    status: str = "init" # init, research_done, writing
    current_chapter_index: int = 0
    context_accumulator: str = ""
    text_parts: List[str] = field(default_factory=list)
    chapter_summaries: List[str] = field(default_factory=list)
    bibliography: List[Dict] = field(default_factory=list)
    
    def to_json(self):
        return json.dumps(asdict(self), ensure_ascii=False, indent=2)
    
    @staticmethod
    def from_json(json_str):
        data = json.loads(json_str)
        return MissionState(**data)

# ==============================================================================
# 2. МОДЕЛИ ДАННЫХ (PYDANTIC) - СТРУКТУРА ЗНАНИЙ 
# ==============================================================================

class ComplexityLevel(str, Enum):
    """Уровень когнитивной нагрузки задачи."""
    ROUTINE = "routine"       # Классификация, форматирование, простой перевод
    ANALYTICAL = "analytical" # Планирование, поиск противоречий, синтез
    CREATIVE = "creative"     # Написание текста, генерация метафор

class Citation(BaseModel):
    """Модель цитаты или источника."""
    source_id: str = Field(..., description="Уникальный ID источника (URL, DOI, Arxiv ID).")
    text_snippet: str = Field(..., description="Фрагмент текста, подтверждающий факт.")
    relevance_score: float = Field(0.0, description="Оценка релевантности (0-1).")

class ChapterBlueprint(BaseModel):
    """
    Детальное техническое задание для одной главы.
    """
    title: str = Field(..., description="Интригующее и содержательное название главы.")
    purpose: str = Field(default="Раскрытие темы", description="Роль главы...")
    core_thesis: str = Field(default="Раскрыть ключевые аспекты раздела.", description="Основной тезис...")
    
    # Модель может вернуть список или строку, валидатор ниже это исправит
    key_points: List[str] = Field(default_factory=lambda: ["Анализ", "Факторы", "Выводы"], description="Список подтем.")
    
    # Здесь была ошибка: модель возвращала список вместо строки
    narrative_link: str = Field(default="Переход к следующей части.", description="Связь со следующей главой.")

    @field_validator('key_points', mode='before')
    def parse_string_list(cls, v):
        if isinstance(v, str):
            return [x.strip() for x in v.split(',') if x.strip()]
        return v

    # === НОВЫЙ ВАЛИДАТОР ===
    @field_validator('narrative_link', 'purpose', 'core_thesis', mode='before')
    def parse_list_to_string(cls, v):
        """Если модель вернула список ['Текст'], превращаем его в строку 'Текст'."""
        if isinstance(v, list):
            return " ".join(str(x) for x in v)
        return v

class ArticleMasterPlan(BaseModel):
    """
    Комплексный план статьи.
    """
    main_title: str = Field(..., description="Основное название всей статьи/книги.")
    subtitle: Optional[str] = Field(default="", description="Подзаголовок.")
    abstract_objective: str = Field(default="Комплексный анализ темы.", description="Краткая аннотация.")
    chapters: List[ChapterBlueprint] = Field(..., description="Список детальных планов.")

class ReviewFeedback(BaseModel):
    """Модель обратной связи от Агента-Критика (Soft Mode)."""
    score: int = Field(default=5, ge=0, le=10)
    strengths: List[str] = Field(default_factory=list)
    weaknesses: List[str] = Field(default_factory=list)
    required_edits: List[str] = Field(default_factory=list)
    is_approved: bool = Field(default=False)



# ==============================================================================
# 3. МЕНЕДЖЕР ТЕКСТА И ТОКЕНОВ
# ==============================================================================

class TextProcessor:
    """
    Утилита для жесткого контроля контекстного окна.
    Гарантирует, что мы никогда не отправим модели больше, чем она может съесть.
    """
    @staticmethod
    def count_tokens(text: str) -> int:
        # Пессимистичная оценка: 1 токен ~ 2.5 символа для смешанного текста
        return int(len(text) / 2.5)

    @staticmethod
    def smart_trim(text: str, max_tokens: int, keep_start: bool = False) -> str:
        """Обрезает текст, сохраняя смысл."""
        current_tokens = TextProcessor.count_tokens(text)
        if current_tokens <= max_tokens:
            return text
        
        # Если нужно сохранить начало (например, введение), режем конец
        # Если нужно сохранить конец (контекст диалога), режем начало
        chars_to_keep = max_tokens * 3
        
        if keep_start:
            return text[:chars_to_keep] + "\n...[TRUNCATED]..."
        else:
            return "...[TRUNCATED]...\n" + text[-chars_to_keep:]

    @staticmethod
    def summarize_context(llm, text: str, max_summary_tokens: int = 500) -> str:
        """Если текст слишком огромный, сжимаем его через LLM."""
        if TextProcessor.count_tokens(text) < 2000:
            return text # Не сжимаем, если и так мало
            
        prompt = f"Текст:\n{text[:6000]}\n\nЗАДАЧА: Сжать этот текст до {max_summary_tokens} токенов, сохранив все ключевые факты и цифры."
        return llm.generate_text("System: Summarizer", prompt, ComplexityLevel.ROUTINE)


# ==============================================================================
# 4. ПРОМПТ-ИНЖИНИРИНГ (УНИВЕРСАЛЬНАЯ ВЕРСИЯ)
# ==============================================================================

class PromptsLibrary:
    """
    Библиотека v4.0 "Universal Architect".
    Специализация: Генерация глубоких, структурированных аналитических документов (отчеты, статьи, книги).
    Основной принцип: Мышление -> Планирование -> Исследование -> Создание.
    """

    # ---  1: системный мандат ---
    # Мы убрали привязку к конкретному журналу и предметной области.
    # Теперь модель - это эксперт широкого профиля.
    SYSTEM_MANDATE = (
        "РОЛЬ: Вы — Мастер-аналитик, экспертная система, специализирующаяся на создании глубоких, структурированных и основанных на данных аналитических документов.\n"
        "ЦЕЛЬ: Преобразовать сложную тему в ясный, логичный и исчерпывающий текст, который будет полезен для информированной аудитории (руководители, аналитики, исследователи).\n\n"
        "ОСНОВНЫЕ ПРИНЦИПЫ РАБОТЫ:\n"
        "1.  **МЫШЛЕНИЕ ПРЕЖДЕ ВСЕГО:** Прежде чем писать, ты глубоко анализируешь тему. Ты выявляешь ключевые вопросы, основные противоречия и логические связи. Твоя главная задача — создать структуру, которая сама по себе несет ценность.\n\n"
        "2.  **ЯСНОСТЬ И ГЛУБИНА:**\n"
        "    - Используй ясный, точный и авторитетный язык. Избегай как излишнего упрощения, так и неоправданного наукообразия.\n"
        "    - Полный запрет на 'воду' и общие фразы ('в современном мире', 'бурное развитие'). Каждый абзац должен нести конкретную мысль.\n"
        "    - Профессиональную терминологию используй по мере необходимости, но всегда в контексте, понятном для умного читателя.\n\n"
        "3.  **ЛОГИКА ИЗЛОЖЕНИЯ (КОНТЕКСТ -> АНАЛИЗ -> СИНТЕЗ -> ВЫВОДЫ):**\n"
        "    - **Контекст:** Всегда начинай с определения проблемы, ее актуальности и необходимых вводных данных.\n"
        "    - **Анализ:** Детально разбирай тему на составные части. Исследуй причины, факторы, механизмы.\n"
        "    - **Синтез:** Собирай разрозненные данные в единую картину. Показывай, как части влияют друг на друга. Предлагай решения, модели или прогнозы.\n"
        "    - **Выводы:** Формулируй четкие, практически применимые выводы на основе проведенного анализа.\n\n"
        "4.  **ДАННЫЕ КАК ОСНОВА:**\n"
        "    - Любое утверждение должно быть подкреплено логикой или фактами, полученными на этапе исследования.\n"
        "    - **КРИТИЧЕСКИ ВАЖНО:** Ты должен органично вплетать найденную информацию (статистику, примеры, цитаты) в свой текст. **ЗАПРЕЩЕНО использовать сноски и ссылки в формате [1], [2] в теле текста.** Информация должна стать частью повествования, а не отдельной сущностью. Список источников будет сформирован в конце документа отдельно."
    )

    # ---  2: промпт Архитектора с фазой "Когнитивного сканирования" ---
    # Это реализует твое требование о "очень, очень объёмном плане".
    ARCHITECT_ROLE = (
        f"{SYSTEM_MANDATE}\n\n"
        "ПРОТОКОЛ: СТРАТЕГИЧЕСКОЕ ПРОЕКТИРОВАНИЕ ЗНАНИЙ.\n"
        "ТВОЯ РОЛЬ: Ты — Архитектор Смыслов. Твоя задача — не просто набросать оглавление, а создать фундаментальный чертеж (Master Plan) для будущего документа.\n"
        "ЗАДАЧА: Разработать исчерпывающий, логически безупречный и детализированный план для комплексного раскрытия заданной темы.\n\n"
        "ПРОЦЕСС РАБОТЫ В ДВА ЭТАПА:\n"
        "ЭТАП 1: КОГНИТИВНОЕ СКАНИРОВАНИЕ (Внутренний монолог, не для вывода)\n"
        "   - *Какие фундаментальные вопросы лежат в основе этой темы?*\n"
        "   - *Какова конечная цель документа? Что должен понять читатель?*\n"
        "   - *Каков наилучший путь для читателя? От простого к сложному? От общего к частному?*\n"
        "   - *Какие подтемы являются ключевыми, а какие — второстепенными?*\n"
        "   - *Существуют ли в теме противоречия или разные точки зрения, которые нужно осветить?*\n\n"
        "ЭТАП 2: ФОРМИРОВАНИЕ MASTER PLAN (Вывод в формате JSON)\n"
        "   - На основе результатов сканирования, создай структуру документа.\n"
        "   - Используй универсальную логическую последовательность (если тема не диктует иного):\n"
        "     1. **ВВЕДЕНИЕ И ПОСТАНОВКА ПРОБЛЕМЫ:** (Почему это важно? Какова цель?)\n"
        "     2. **ТЕОРЕТИЧЕСКИЙ ФУНДАМЕНТ:** (Основные понятия, принципы, история вопроса).\n"
        "     3. **КЛЮЧЕВОЙ АНАЛИЗ:** (Разбор основной части темы, исследование факторов, механизмов).\n"
        "     4. **ПРАКТИЧЕСКОЕ ПРИМЕНЕНИЕ / КЕЙСЫ:** (Как это работает в реальном мире? Примеры, данные).\n"
        "     5. **СИНТЕЗ И ПЕРСПЕКТИВЫ:** (Объединение всех данных, выводы, прогнозы, рекомендации).\n\n"
        "ТРЕБОВАНИЕ К JSON ВЫВОДУ:\n"
        "Для каждой главы в поле `key_points` перечисли КОНКРЕТНЫЕ вопросы или аспекты, которые должны быть раскрыты (например: 'Анализ SWOT для технологии X', 'Сравнение подходов A и B', 'Разбор кейса компании Y')."
    )

    # --- 3: Промпты поиска ---
    QUERY_GENERATOR_PROTOCOL = (
        f"{SYSTEM_MANDATE}\n\n"
        "ПРОТОКОЛ: СТРАТЕГИЧЕСКИЙ ПОИСК ИНФОРМАЦИИ (STRATEGIC OSINT).\n"
        "ТВОЯ РОЛЬ: Аналитик-исследователь. Твоя задача — сформулировать запросы для поиска фактов.\n\n"
        "СТРАТЕГИЯ ПОИСКА ЗАВИСИТ ОТ ЦЕЛИ ГЛАВЫ:\n"
        "1.  **ДЛЯ ВВЕДЕНИЯ И ОПИСАНИЯ ПРОБЛЕМЫ (Поиск Контекста):**\n"
        "    - Ищи статистику, рыночные отчеты, исторические данные.\n"
        "    - Формат: '[ГЛАВНАЯ ТЕМА] статистика 2025', 'рынок [ГЛАВНАЯ ТЕМА] объем'.\n\n"
        "2.  **ДЛЯ АНАЛИЗА И МЕТОДОЛОГИИ (Поиск Компонентов):**\n"
        "    - **ВАЖНО:** Не ищи общие понятия в отрыве от темы.\n"
        "    - **ПРИМЕР:** Если тема 'Кофе', а глава про 'Логистику', ищи 'Логистика КОФЕ', а не просто 'Логистика'.\n"
        "    - Ищи технические стандарты и методы, специфичные именно для НАШЕЙ ТЕМЫ.\n\n"
        "3.  **ДЛЯ ПРИМЕРОВ И СРАВНЕНИЯ (Поиск Прецедентов):**\n"
        "    - Ищи кейсы и отчеты о внедрении ИМЕННО В ЭТОЙ ИНДУСТРИИ.\n"
        "    - Формат: '[ГЛАВНАЯ ТЕМА] best practices', 'кейс [ПОДТЕМА] в сфере [ГЛАВНАЯ ТЕМА]'.\n\n"
        "ВЫВОД: JSON с 1 запросам. ВСЕ ЗАПРОСЫ ДОЛЖНЫ СОДЕРЖАТЬ КЛЮЧЕВОЕ СЛОВО ТЕМЫ."
    )

    # --- 4: Досье  ---
    EVIDENCE_DISTILLER_PROTOCOL = (
        f"{SYSTEM_MANDATE}\n\n"
        "ПРОТОКОЛ: ЭКСТРАКЦИЯ ФАКТОВ ИЗ ИСТОЧНИКА.\n"
        "ТВОЯ РОЛЬ: Специалист по анализу данных. Тебе на вход подан текст. Твоя задача — отфильтровать 'воду' и субъективные мнения, оставив только сухие, проверяемые факты для нашего документа.\n\n"
        "INPUT:\n"
        "1.  **Контекст:** Тема главы (что мы ищем).\n"
        "2.  **Источник:** Текст для анализа.\n\n"
        "МЕТОДОЛОГИЯ:\n"
        "1.  **ФИЛЬТР РЕЛЕВАНТНОСТИ:** Если текст не относится к теме — верни '[RELEVANCE: LOW]'.\n\n"
        "2.  **ЭКСТРАКЦИЯ ДАННЫХ:** Если источник полезен, заполни АНАЛИТИЧЕСКОЕ ДОСЬЕ. Будь предельно конкретен. Копируй данные точно.\n\n"
        "   --- АНАЛИТИЧЕСКОЕ ДОСЬЕ ---\n\n"
        "   - **КОЛИЧЕСТВЕННЫЕ ДАННЫЕ (Metrics & Data):**\n"
        "     [Выпиши ВСЕ цифры: проценты, суммы, даты, статистика, размеры выборок, любые измеримые показатели.]\n\n"
        "   - **КЛЮЧЕВЫЕ СУЩНОСТИ (Entities & Concepts):**\n"
        "     [Выпиши названия компаний, продуктов, технологий, стандартов, имена экспертов, ключевые термины и их определения.]\n\n"
        "   - **ОСНОВНЫЕ ТЕЗИСЫ И ВЫВОДЫ (Arguments & Conclusions):**\n"
        "     [Кратко, своими словами или цитатами: Какова главная мысль автора? Какие проблемы он выделяет? Какие решения предлагает?]\n\n"
        "   --- КОНЕЦ ДОСЬЕ ---\n\n"
        "3.  **ВЫВОД:** Только заполненное досье или метка LOW."
    )

    WRITER_PHASE_1_ANALYSIS = (
        f"{SYSTEM_MANDATE}\n\n"
        "ПРОТОКОЛ: АНАЛИТИЧЕСКИЙ БРИФ.\n"
        "ЦЕЛЬ: Продумать внутреннюю логику и структуру главы перед написанием.\n"
        "ДЕЙСТВИЯ:\n"
        "1. Определи главный тезис, который ты докажешь в этой главе.\n"
        "2. Составь пошаговый план изложения внутри главы (микро-структура).\n"
        "3. Какие данные из досье будут наиболее убедительны для каждого шага?\n"
        "ВЫВОД: Краткий аналитический бриф."
    )

    WRITER_PHASE_2_EXECUTION = (
        f"{SYSTEM_MANDATE}\n\n"
        "ПРОТОКОЛ: СОЗДАНИЕ АВТОРСКОГО ТЕКСТА.\n\n"
        "ПРАВИЛА:\n"
        "1.  **ПОВЕСТВОВАНИЕ, А НЕ ОТЧЕТ:** Твоя задача — не перечислить факты, а выстроить на их основе убедительное повествование. Каждый факт должен работать на общую идею главы.\n\n"
        "2.  **ФАКТИЧЕСКОЕ ЗАЗЕМЛЕНИЕ (CRITICAL):** ЗАПРЕЩЕНО приводить конкретные числовые данные (проценты, суммы, даты, статистика), если их нет в предоставленном блоке 'ВНЕШНИЕ ФАКТЫ'. Если точных цифр нет, используй общие формулировки ('значительно выросли', 'существенная часть расходов', 'в последние годы'). Не выдумывай статистику для убедительности.\n\n"
        "3.  **ИНТЕГРАЦИЯ ДАННЫХ:** Органично вплетай найденные данные (цифры, примеры, факты) в текст. Например, вместо 'Точность системы составила 95% [3]', напиши 'Благодаря новому подходу, точность системы достигла впечатляющих 95%, что позволило...'.\n\n"
        "4.  **СТРУКТУРА И ЧИТАЕМОСТЬ:** Используй абзацы для разделения мыслей. Если нужно, используй Markdown для списков и таблиц, чтобы сделать сложную информацию наглядной.\n\n"
        "5.  **ЛОГИЧЕСКИЕ ПЕРЕХОДЫ:** Обеспечь плавные переходы между абзацами и подтемами. Читатель должен легко следовать за твоей мыслью.\n\n"
        "6.  **ЗАПРЕТ НА ССЫЛКИ:** Еще раз: **НЕ ИСПОЛЬЗУЙ ССЫЛКИ-СНОСКИ формата [1], [2], [источник] в тексте.**\n\n"
        "ВЫВОД: Полный, хорошо написанный и аргументированный текст главы."
    )

    STYLE_POLYMATH_DIRECTIVE = (
        f"{SYSTEM_MANDATE}\n"
        "ЗАДАЧА: Стилистическое обогащение текста.\n"
        "1. Замени повторяющиеся слова и конструкции на синонимы.\n"
        "2. Сделай слишком сложные предложения проще, а слишком простые — объедини в более комплексные для динамики.\n"
        "3. Проверь, чтобы тон повествования был уверенным и экспертным.\n"
        "4. Убери канцеляризмы и штампы.\n"
        "ВЫВОД: Отредактированный, стилистически выверенный текст."
    )

    CONTINUITY_DIRECTIVE = (
        f"{SYSTEM_MANDATE}\n"
        "ЗАДАЧА: Обеспечение логической связности.\n"
        "Проанализируй вывод предыдущей главы и цель текущей. Напиши 'мостик' — один-два абзаца, которые плавно и логично свяжут эти две части, создавая единое повествование."
    )
    
    CRITIC_DIRECTIVE = (
        f"{SYSTEM_MANDATE}\n"
        "РОЛЬ: Внутренний критик.\n"
        "1. Достаточно ли сильна аргументация? Все ли тезисы подкреплены фактами или логикой?\n"
        "2. Нет ли в тексте логических противоречий или необоснованных обобщений?\n"
        "3. Понятен ли текст для целевой аудитории? Не упущено ли что-то важное?\n"
        "ВЫВОД: JSON с оценкой и конкретными рекомендациями по улучшению."
    )

    # Промпты для сборки и публикации 
    FINAL_ASSEMBLER_PROTOCOL = (
        "РОЛЬ: Выпускающий редактор. Твоя задача — собрать из разрозненных глав единый, чистый и логически выстроенный манускрипт.\n"
        "ЦЕЛЬ: Устранить дублирование, артефакты генерации и обеспечить идеальную структуру документа перед финальной сдачей.\n\n"
        "ПОШАГОВЫЙ ПРОЦЕСС РЕДАКТУРЫ:\n"
        "1.  **СТРУКТУРНАЯ СБОРКА:** Расположи все части в правильном порядке: Заголовок -> Аннотация -> Ключевые слова -> Введение -> Основные главы -> Заключение -> Список литературы.\n"
        "2.  **УДАЛЕНИЕ ДУБЛИКАТОВ:** В тексте могут быть дублирующиеся заголовки или разделы (например, несколько 'Заключений'). Оставь только один, самый полный и релевантный вариант. Остальные удали.\n"
        "3.  **ОЧИСТКА ОТ АРТЕФАКТОВ:** Удали все служебные пометки, мета-комментарии, фразы типа 'Вот текст:' и прочий технический мусор, который мог остаться от предыдущих агентов.\n"
        "4.  **СГЛАЖИВАНИЕ ПЕРЕХОДОВ:** Проверь переходы между главами. Если они слишком резкие, добавь одно-два связующих предложения для плавности.\n\n"
        "ВЫВОД: Только идеально чистый и отформатированный текст, готовый к публикации. Без каких-либо твоих комментариев."
    )

    PUBLISHER_DIRECTIVE = (
        f"{SYSTEM_MANDATE}\n\n"
        "РЕЖИМ: АБСТРАКТНЫЙ СИНТЕЗ И ФИНАЛИЗАЦИЯ.\n"
        "ВХОДНЫЕ ДАННЫЕ: Сводка (Summary) содержания всех глав документа.\n\n"
        "ЗАДАЧА: Сгенерировать три финальных служебных раздела.\n\n"
        "*** СТРОГИЕ ПРАВИЛА ГЕНЕРАЦИИ ***\n"
        "1.  **ПРИНЦИП «ЗЕРКАЛА»:** Твой ответ должен отражать ИСКЛЮЧИТЕЛЬНО те темы, факты и выводы, которые есть во входном тексте. Запрещено добавлять информацию «от себя» или использовать шаблонные фразы про «современный мир», «инновации» или «технологии», если этих тем нет в источнике.\n"
        "2.  **ФОРМАТ ЗАГОЛОВКОВ:** Используй Markdown (# Заголовок). КАТЕГОРИЧЕСКИ ЗАПРЕЩЕНО ставить цифры или номера перед заголовками.\n\n"
        "ТРЕБУЕМАЯ СТРУКТУРА ОТВЕТА:\n\n"
        "# АННОТАЦИЯ\n"
        "(Сжатая суть документа в 1 абзац. О чем этот текст? Какова его главная мысль? Пиши сухо и по делу).\n\n"
        "# КЛЮЧЕВЫЕ СЛОВА\n"
        "(Список из 10-15 главных терминов, которые чаще всего встречаются в тексте, через запятую).\n\n"
        "# ЗАКЛЮЧЕНИЕ\n"
        "(Финальный итог работы. Не пересказывай главы, а сформулируй общий вывод, вытекающий из всего текста.)"
    )
    
    # Системные утилиты остаются без изменений.
    CONTEXT_SUMMARY_PROTOCOL = (
        f"{SYSTEM_MANDATE}\n"
        "ЗАДАЧА: Сжать предоставленный текст до 2-3 ключевых тезисов. Эта информация будет использована как контекстная память для следующих этапов. Сохрани все важные цифры и выводы."
    )
    
    GHOST_PROTOCOL = "System: Ensure maximum logical density and clarity."
    JSON_REPAIR_UTILITY = "System: The provided text is a broken JSON. Your only task is to fix syntax errors (missing brackets, commas, quotes) and return a valid JSON object. Do not change the data."

# ==============================================================================
# 5. КОГНИТИВНЫЙ ДВИЖОК (LLM ENGINE) - СЕРДЦЕ СИСТЕМЫ 
# ==============================================================================

class LLMEngine:
    """
    Обертка над OpenAI Client.
    Версия v4.0: ROBUST JSON REPAIR + CONTEXT AWARENESS.
    Исправлена ошибка "Амнезии" при починке JSON.
    """
    def __init__(self, config: SystemConfig):
        self.config = config
        self.client = None
        
        self.log_dir = "genesis_thoughts"
        os.makedirs(self.log_dir, exist_ok=True)
        self.live_log_file = os.path.join(
            self.log_dir, 
            f"LIVE_THINK_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        )
        
        with open(self.live_log_file, 'w', encoding='utf-8') as f:
            # Исправлено название на актуальное
            f.write(f"=== GENESIS GENERATOR LIVE THOUGHT STREAM (SANITIZED) ===\n")
            f.write(f"Started at: {datetime.now()}\n\n")
            
        self._connect_to_server()

    def _connect_to_server(self):
        """
        Инициализация соединения с проверкой доступности.
        """
        base_url = self.config.LLM_BASE_URL
        
        # Пытаемся взять ключ из конфига, если его нет — ставим "lm-studio" по умолчанию
        api_key = getattr(self.config, "LLM_API_KEY", "lm-studio") 
        
        logger.info("Проверка соединения с LM Studio...")
        
        try:
            # 1. Физическая проверка доступности порта
            # Пытаемся получить список моделей (стандартный эндпоинт OpenAI)
            test_url = f"{base_url}/models"
            response = requests.get(test_url, timeout=3)
            
            if response.status_code == 200:
                models = response.json().get('data', [])
                model_id = models[0]['id'] if models else "unknown"
                logger.info(f"✅ Сервер доступен. Активная модель: {model_id}")
            else:
                logger.warning(f"⚠️ Сервер ответил кодом {response.status_code}. Проверьте LM Studio.")

            # 2. Инициализация клиента
            self.client = OpenAI(base_url=base_url, api_key=api_key)
            logger.info(f"OpenAI Client инициализирован. Живой лог: {self.live_log_file}")

        except requests.exceptions.ConnectionError:
            logger.critical("\n❌ ОШИБКА ПОДКЛЮЧЕНИЯ: Не удалось связаться с LM Studio.")
            logger.critical(f"Убедитесь, что сервер запущен на {base_url}")
            logger.critical("Запустите LM Studio -> Start Server.\n")
            # Не выходим через sys.exit, чтобы GUI мог показать ошибку, но ставим флаг
            self.client = None
        except Exception as e:
            logger.critical(f"❌ Непредвиденная ошибка инициализации: {e}")
            self.client = None

    def _get_model_id(self, complexity: ComplexityLevel) -> str:
        return self.config.MODEL_SMART_ID if complexity != ComplexityLevel.ROUTINE else self.config.MODEL_FAST_ID

    def _clean_llm_output(self, text: str) -> str:
        if not text: return ""
        
        # 1. Удаляем "мысли" (DeepSeek)
        # !!! ВАЖНО: Заменяем на ПРОБЕЛ (' '), а не на пустоту (''), чтобы слова не склеивались
        text = re.sub(r'<think>.*?</think>', ' ', text, flags=re.DOTALL | re.IGNORECASE)
        
        # 2. Удаляем Markdown блоки кода
        text = re.sub(r'^```[a-zA-Z]*\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r'\s*```$', '', text, flags=re.MULTILINE)
        
        # 3. Удаляем технические заголовки
        # Используем MULTILINE вместо DOTALL, чтобы удалять только мусорные строки,
        # а не случайно снести половину текста.
        patterns_to_kill = [
            r'^PHASE \d+:.*$', r'^Фаза \d+:.*$', r'^STAGE \d+:.*$',
            r'^Here is the (JSON|text|output).*?[:\n]', 
            r'^Вот (текст|JSON|результат).*?[:\n]', 
            r'^Based on.*?[:\n]', 
            r'^Sure,.*?[:\n]', r'^Certainly.*?[:\n]', r'^Конечно.*?[:\n]',
            r'\[RELEVANCE:.*?\]' # Удаляем метки релевантности, если остались
        ]
        
        for pattern in patterns_to_kill:
            # Заменяем на пустоту, так как удаляем целые строки мусора
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
            
        # Чистим двойные пробелы, возникшие после удалений
        text = re.sub(r'[ \t]+', ' ', text)
        # Чистим множественные переносы строк (оставляем максимум 2 для абзацев)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        text = text.strip()

        # 4. ФИКС ОБРЫВОВ (Cut-off Fix)
        valid_endings = ['.', '!', '?', '"', '»', '…', ')', ']', '}', ';']
        
        # Проверяем длину (>50), чтобы не убить короткие ответы типа "Да."
        if len(text) > 50 and text[-1] not in valid_endings:
            last_punctuation = -1
            for char in ['.', '!', '?', '…']:
                pos = text.rfind(char)
                if pos > last_punctuation:
                    last_punctuation = pos
            
            # Обрезаем, только если сохраним хотя бы 20% текста
            if last_punctuation > len(text) * 0.2:
                text = text[:last_punctuation+1]
            
        return text

    def generate_text(self, system_prompt: str, user_prompt: str, complexity: ComplexityLevel = ComplexityLevel.ANALYTICAL, temperature: float = 0.7) -> str:
        if not self.client: return "Error: Client not connected"
        
        model_id = self._get_model_id(complexity)
        
        # ===  AUTO-TRIM (УМНЫЙ ПРЕДОХРАНИТЕЛЬ) ===
        # 1. Берем общий лимит памяти из конфига (например, 16384)
        total_window = self.config.CONTEXT_WINDOW_SIZE
        
        # 2. Вычитаем место, зарезервированное под ответ модели (например, 4096)
        # И вычитаем токены системного промпта (он неприкасаем)
        system_tokens = TextProcessor.count_tokens(system_prompt)
        reserved_tokens = self.config.MAX_OUTPUT_TOKENS + system_tokens + 200 # +200 буфер безопасности
        
        # 3. Сколько места осталось для текста пользователя?
        allowed_user_tokens = total_window - reserved_tokens
        
        # 4. Проверяем текущий текст
        current_tokens = TextProcessor.count_tokens(user_prompt)
        
        if current_tokens > allowed_user_tokens:
            # Если перебор — режем!
            logger.warning(f" ПЕРЕГРУЗКА ПАМЯТИ: {current_tokens} токенов. Лимит: {allowed_user_tokens}.")
            logger.warning("   -> Удаляю старый контекст (начало), сохраняю задачу (конец)...")
            
            # keep_start=False -> Удаляет начало текста (старое), сохраняет конец (свежее)
            user_prompt = TextProcessor.smart_trim(user_prompt, allowed_user_tokens, keep_start=False)
        # ============================================

        messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]

        for attempt in range(self.config.MAX_LLM_RETRIES):
            try:
                with open(self.live_log_file, 'a', encoding='utf-8') as f:
                    # Логируем только конец промпта, чтобы не засорять файл, если он обрезан
                    log_prompt = user_prompt[-500:] if len(user_prompt) > 500 else user_prompt
                    f.write(f"\n{'='*60}\n {datetime.now().strftime('%H:%M:%S')} | MODEL: {model_id}\nTASK (Last 500 chars): ...{log_prompt}\n{'='*60}\n [STREAM]:\n")
                    f.flush()

                    stream = self.client.chat.completions.create(
                        model=model_id,
                        messages=messages,
                        temperature=temperature,
                        max_tokens=self.config.MAX_OUTPUT_TOKENS,
                        stream=True
                    )
                    
                    full_response = ""
                    reasoning_buffer = ""
                    
                    for chunk in stream:
                        delta = chunk.choices[0].delta
                        content = getattr(delta, "content", "") or ""
                        reasoning = getattr(delta, "reasoning_content", "") or ""

                        if reasoning:
                            reasoning_buffer += reasoning

                        if content:
                            if reasoning_buffer:
                                thought_block = f"<think>\n{reasoning_buffer.strip()}\n</think>\n"
                                f.write(thought_block)
                                full_response += thought_block
                                reasoning_buffer = "" 

                            f.write(content)
                            full_response += content
                        
                        f.flush() 

                    if reasoning_buffer:
                        thought_block = f"<think>\n{reasoning_buffer.strip()}\n</think>\n"
                        f.write(thought_block)
                        full_response += thought_block

                    f.write(f"\n\n[END]\n")
                
                # Телеметрия (приблизительная)
                t_in = len(user_prompt) // 3
                t_out = len(full_response) // 3
                logger.info(f"[TELEMETRY] {{'in': {t_in}, 'out': {t_out}, 'req': 1}}")
                
                return self._clean_llm_output(full_response)

            except Exception as e:
                logger.warning(f"LLM Error ({attempt+1}): {e}")
                
                # Если ошибка всё равно вылезла (иногда модель считает токены иначе), режем еще агрессивнее
                if "context" in str(e).lower() or "length" in str(e).lower():
                    logger.error("Повторный сбой контекста. Режу еще 30% текста...")
                    user_prompt = user_prompt[int(len(user_prompt)*0.3):] # Отрезаем первую треть
                
                time.sleep(1)
        return ""

    def generate_structured(self,
                            system_prompt: str,
                            user_prompt: str,
                            response_model: Type[BaseModel],
                            complexity: ComplexityLevel = ComplexityLevel.ANALYTICAL,
                            temperature: float = 0.7) -> BaseModel:
        """
        Метод для получения строгого JSON с многоуровневой защитой v5.0.
        Совмещает надежность старой версии с гибкостью новой.
        """
        # Генерация примера JSON для промпта (улучшенная версия)
        example_json = "{}"
        model_name = response_model.__name__
        if model_name == "ArticleMasterPlan":
            example_json = '{"main_title": "Заголовок", "chapters": [{"title": "Глава 1", "purpose": "Цель"}, ...]}'
        elif model_name == "SearchQueryPlan":
            example_json = '{"queries": ["поисковый запрос 1", "поисковый запрос 2"]}'
        elif model_name == "ReviewFeedback":
            example_json = '{"score": 8, "is_approved": true, "required_edits": ["Исправить ..."]}'

        last_error = None
        
        # === УРОВЕНЬ 1: СТАНДАРТНЫЙ JSON (с поддержкой списков) ===
        for attempt in range(2):
            try:
                augmented_prompt = (
                    f"{user_prompt}\n\n"
                    f"*** СИСТЕМА: ВЕРНИ ТОЛЬКО RAW VALID JSON ОБЪЕКТ ***\n"
                    f"Пример формата:\n{example_json}"
                )
                
                raw_response = self.generate_text(system_prompt, augmented_prompt, complexity, temperature)
                
                json_str = self._extract_json_string(raw_response)
                if json_str: json_str = self._auto_close_json(json_str)
                if not json_str: raise ValueError("JSON braces or brackets not found in the response")
                
                # Умный парсинг: если модель вернула чистый список, оборачиваем его в объект
                if json_str.strip().startswith('['):
                    field_name = list(response_model.model_fields.keys())[0]
                    data_dict = {field_name: json.loads(json_str)}
                else:
                    data_dict = json.loads(json_str)

                # Используем твою утилиту для "разворачивания" вложенного JSON
                data_dict = self._unwrap_nested_json(data_dict, response_model)
                return response_model.model_validate(data_dict)

            except Exception as e:
                last_error = e
                logger.warning(f"JSON Level 1 (Attempt {attempt+1}) failed for {model_name}: {e}")

        # === УРОВЕНЬ 2: АДАПТИВНЫЙ ТЕКСТОВЫЙ ФОРМАТ (План Б) ===
        try:
            logger.info(f"Switching to Level 2 (Simple Format) for {model_name}...")
            
            if model_name == "ArticleMasterPlan":
                simple_prompt = (
                    f"{user_prompt}\n\n"
                    f"*** СИСТЕМА: JSON СЛОМАЛСЯ. ИСПОЛЬЗУЙ ПРОСТОЙ ФОРМАТ. ***\n"
                    f"Не используй скобки или запятые. Просто пиши:\n"
                    f"TITLE: <название статьи>\n"
                    f"SUBTITLE: <подзаголовок>\n"
                    f"CHAPTER 1: <название главы 1>\n"
                    f"CHAPTER 2: <название главы 2>\n..."
                )
                raw_text = self.generate_text(system_prompt, simple_prompt, complexity, temperature)
                return self._parse_simple_master_plan(raw_text)
            
            elif model_name == "SearchQueryPlan":
                simple_prompt = (
                    f"{user_prompt}\n\n"
                    f"*** СИСТЕМА: JSON СЛОМАЛСЯ. ИСПОЛЬЗУЙ ПРОСТОЙ ФОРМАТ. ***\n"
                    f"Не используй скобки или запятые. Просто перечисли поисковые запросы, каждый на новой строке, начиная с тире '- '."
                )
                raw_text = self.generate_text(system_prompt, simple_prompt, complexity, temperature)
                queries = [line.strip('- ') for line in raw_text.split('\n') if line.strip()]
                return response_model(queries=queries)

            # Если для модели нет "Плана Б", сразу переходим к фоллбэку
            raise NotImplementedError(f"No Level 2 parser available for {model_name}")

        except Exception as e:
             logger.error(f"Level 2 failed for {model_name}: {e}")

        # === УРОВЕНЬ 3: АВАРИЙНЫЙ ФОЛЛБЭК ===
        logger.error(f"FATAL: All generation methods failed for {model_name}. Last error: {last_error}")
        return self._create_fallback_object(response_model, user_prompt)

    def _parse_simple_master_plan(self, text: str) -> BaseModel:
        """
        Усиленный парсер: вытаскивает структуру из обычного текста, 
        если JSON не сработал. Спасает реальные идеи модели.
        """
        main_title = "Draft Plan"
        subtitle = ""
        abstract = "Анализ темы"
        chapters = []
        
        # 1. Попытка вытащить Заголовок и Абстракт через регулярки
        title_match = re.search(r'(?:Title|Заголовок|Тема):\s*(.+)', text, re.IGNORECASE)
        if title_match: main_title = title_match.group(1).strip(' *"«»')
        
        abs_match = re.search(r'(?:Abstract|Аннотация|Цель):\s*(.+)', text, re.IGNORECASE)
        if abs_match: abstract = abs_match.group(1).strip()

        # 2. Разбиваем текст на блоки по главам
        # Ищем строки, начинающиеся с "Глава X" или "Chapter X" или просто "1."
        lines = text.split('\n')
        current_chapter = None
        
        for line in lines:
            line = line.strip()
            if not line: continue
            
            # Эвристика начала главы
            is_chapter_start = re.match(r'^(?:Chapter|Глава|\d+\.)\s', line, re.IGNORECASE)
            
            if is_chapter_start:
                # Если уже собирали главу, сохраняем её
                if current_chapter:
                    chapters.append(current_chapter)
                
                # Начинаем новую
                # Очищаем заголовок от "Глава 1: "
                clean_title = re.sub(r'^(?:Chapter|Глава|\d+)\s*[:.]\s*', '', line, flags=re.IGNORECASE).strip()
                if "|" in clean_title: clean_title = clean_title.split('|')[0].strip()

                current_chapter = ChapterBlueprint(
                    title=clean_title,
                    purpose="Section",
                    core_thesis=f"Раскрыть тему: {clean_title}", # Временный тезис из названия
                    key_points=[],
                    narrative_link="Далее"
                )
            
            # Если мы внутри главы, пытаемся вытащить детали
            elif current_chapter:
                # Если строка похожа на пункт списка (- или *)
                if line.startswith('-') or line.startswith('*'):
                    point = line.strip('-* ').strip()
                    if len(point) > 5:
                        current_chapter.key_points.append(point)
                
                # Если строка похожа на тезис/цель
                elif "Goal:" in line or "Цель:" in line or "Тезис:" in line:
                    thesis = line.split(':', 1)[1].strip()
                    current_chapter.core_thesis = thesis

        # Добавляем последнюю главу
        if current_chapter:
            chapters.append(current_chapter)

        if not chapters and len(text) > 50:
             # Просто создаем одну главу из всего текста
             chapters.append(ChapterBlueprint(
                 title="Основной анализ", 
                 core_thesis="Рассмотреть предоставленную информацию",
                 key_points=["Анализ данных"]
             ))

        # Возвращаем объект (используем Smart Defaults из шага 1)
        return ArticleMasterPlan(
            main_title=main_title, 
            subtitle=subtitle, 
            abstract_objective=abstract, 
            chapters=chapters
        )

    def _extract_json_string(self, text: str) -> Optional[str]:
        """
        Усиленный поиск JSON/JSON-списка в ответе.
        """
        text = text.strip()
        
        # Попытка 1: Найти блок ```json ... ``` (Markdown)
        match = re.search(r'```json\s*([\[\{].*?[\]\}])\s*```', text, re.DOTALL)
        if match: return match.group(1)
        
        # Попытка 2: Найти блок ``` ... ``` (любой код)
        match = re.search(r'```\s*([\[\{].*?[\]\}])\s*```', text, re.DOTALL)
        if match: return match.group(1)

        # Попытка 3: Найти самые внешние скобки (фигурные или квадратные)
        start_brace = text.find('{')
        start_bracket = text.find('[')

        # Определяем, что начинается раньше
        if start_brace == -1: start = start_bracket
        elif start_bracket == -1: start = start_brace
        else: start = min(start_brace, start_bracket)
        
        if start == -1: return None

        # Определяем, какой символ закрывающей скобки искать
        end_char = '}' if text[start] == '{' else ']'
        
        end = text.rfind(end_char)
        if end > start:
            return text[start : end + 1]
        
        return text[start:] # Для авто-починки

    def _auto_close_json(self, json_str: str) -> str:
        """
        Эвристика для починки оборванного JSON v2.0.
        Исправляет незакрытые кавычки перед закрытием скобок.
        """
        json_str = json_str.strip()
        
        # 1. Проверка на незакрытую строку
        # Считаем количество кавычек (неэкранированных)
        quote_count = 0
        escaped = False
        for char in json_str:
            if char == '\\':
                escaped = not escaped
            elif char == '"' and not escaped:
                quote_count += 1
                escaped = False
            else:
                escaped = False
        
        # Если количество кавычек нечетное, значит строка оборвалась
        if quote_count % 2 != 0:
            json_str += '"' # Закрываем строку принудительно

        # 2. Убираем висячую запятую (частая ошибка)
        if json_str.endswith(','): 
            json_str = json_str[:-1]
        
        # 3. Балансировка скобок
        open_braces = json_str.count('{')
        close_braces = json_str.count('}')
        open_brackets = json_str.count('[')
        close_brackets = json_str.count(']')
        
        # Закрываем в порядке обратной вложенности (упрощенно: сначала ], потом })
        if open_brackets > close_brackets:
            json_str += ']' * (open_brackets - close_brackets)
        
        if open_braces > close_braces:
            json_str += '}' * (open_braces - close_braces)
            
        return json_str

    def _unwrap_nested_json(self, data: Any, model: Type[BaseModel]) -> Dict:
        if not isinstance(data, dict): return data
        req = set(model.model_fields.keys())
        # Если ключи модели есть в корне - ок
        if len(req.intersection(data.keys())) >= 1: return data
        # Ищем вложенный словарь
        for v in data.values():
            if isinstance(v, dict) and len(req.intersection(v.keys())) >= 1: return v
        return data

    def _create_fallback_object(self, model_class: Type[BaseModel], prompt: str) -> BaseModel:
        """
        Создает безопасный аварийный объект, который пытается спасти миссию,
        используя исходный пользовательский промпт. Гарантированно возвращает объект.
        """
        try:
            logger.warning(f"⚠️ JSON Fallback activated for {model_class.__name__}")
            
            # Аварийный план статьи - САМОЕ ВАЖНОЕ
            if model_class.__name__ == "ArticleMasterPlan":
                topic = "Неизвестная тема (сбой извлечения)"
                try:
                    # Улучшенное извлечение темы из любого места промпта
                    match = re.search(r'(?:Тема|Название статьи):\s*["«](.*?)[»"]', prompt, re.IGNORECASE | re.DOTALL)
                    if match:
                        topic = match.group(1).strip()
                except: pass
                
                # ИСПРАВЛЕНИЕ: Используем ChapterBlueprint вместо ArticleSection.
                # Заполняем все поля заглушками, чтобы не было ошибки AttributeError позже.
                return model_class.model_construct(
                    main_title=topic, 
                    subtitle="Аналитический отчет (сгенерирован в аварийном режиме)", 
                    abstract_objective="Автоматическая генерация структуры не удалась. Используется базовый план.", 
                    chapters=[
                        ChapterBlueprint(
                            title="Введение в проблему", 
                            purpose="Introduction",
                            core_thesis="Описать контекст и актуальность проблемы.",
                            key_points=["Определение проблемы", "Текущее состояние", "Цели работы"],
                            narrative_link="Переход к детальному анализу."
                        ),
                        ChapterBlueprint(
                            title="Анализ ключевых аспектов", 
                            purpose="Analysis",
                            core_thesis="Рассмотреть детали и факторы влияния.",
                            key_points=["Фактор 1", "Фактор 2", "Сравнительный анализ"],
                            narrative_link="Переход к практическим решениям."
                        ),
                        ChapterBlueprint(
                            title="Практическое применение", 
                            purpose="Solution",
                            core_thesis="Показать примеры и методологию.",
                            key_points=["Описание метода", "Примеры реализации", "Результаты"],
                            narrative_link="Переход к выводам."
                        ),
                        ChapterBlueprint(
                            title="Выводы и перспективы", 
                            purpose="Conclusion",
                            core_thesis="Подвести итоги и дать прогноз.",
                            key_points=["Основные выводы", "Рекомендации", "Прогноз на будущее"],
                            narrative_link="Финал."
                        )
                    ]
                )
            
            # Другие фоллбэки
            # Обрабатываем и старое название DeepDivePlan, и новое SearchQueryPlan (из ResearcherAgent)
            if "DeepDivePlan" in str(model_class) or "SearchQueryPlan" in str(model_class):
                return model_class.model_construct(queries=[])
            
            if model_class.__name__ == "ReviewFeedback":
                return model_class.model_construct(score=5, is_approved=True, strengths=[], weaknesses=[], required_edits=[])

            # Универсальная заглушка для всего остального
            return model_class.model_construct()
            
        except Exception as e:
             logger.error(f"Fallback creation failed: {e}")
             # Самый крайний случай - возвращаем пустой объект, чтобы не крашить тред
             return model_class.model_construct()

# ==============================================================================
# 6. ВСПОМОГАТЕЛЬНЫЕ УТИЛИТЫ (UTILS) 
# ==============================================================================

def save_text_to_file(content: str, filename: str, subfolder: str = ""):
    """Надежное сохранение текста."""
    path = os.path.join(CONFIG.WORK_DIR, subfolder)
    os.makedirs(path, exist_ok=True)
    full_path = os.path.join(path, filename)
    
    try:
        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(content)
        logger.info(f"Файл сохранен: {full_path}")
        return full_path
    except Exception as e:
        logger.error(f"Ошибка сохранения файла {filename}: {e}")
        return None

def load_text_from_file(filepath: str) -> str:
    """Безопасное чтение текста."""
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            return f.read()
    except Exception as e:
        logger.error(f"Ошибка чтения файла {filepath}: {e}")
        return ""

# ==============================================================================
# 7. КОНФИГУРАЦИЯ СТИЛЯ 
# ==============================================================================
ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("dark-blue")

# Палитра (Neon Cyberpunk)
C_BG = "#050508"         # Void Black
C_PANEL = "#0B0C12"      # Panel Grey
C_BORDER = "#1F222E"     # Border Grey
C_CYAN = "#00F0FF"       # Tech
C_GOLD = "#FFD700"       # Architect
C_BLUE = "#1E90FF"       # Researcher
C_GREEN = "#00FA9A"      # Writer
C_RED = "#FF4500"        # Critic
C_TEXT = "#E0E0E0"

FONT_MONO = ("Consolas", 11)


# ==============================================================================
# 8. ГЛОБАЛЬНЫЕ КОНСТАНТЫ И НАСТРОЙКИ
# ==============================================================================

# Таймауты и ограничения
MAX_DOWNLOAD_SIZE_MB = 25 

# Пути для кеша
CACHE_DIR = os.path.join(CONFIG.WORK_DIR, "system_cache")
os.makedirs(CACHE_DIR, exist_ok=True)

# ==============================================================================
# 9. ВСПОМОГАТЕЛЬНЫЕ МОДУЛИ (UTILS)
# ==============================================================================

class CacheManager:
    """
    Управляет локальным кешированием HTTP-запросов и результатов поиска.
    Позволяет системе работать быстрее и не попадать под бан API.
    """
    def __init__(self, cache_dir: str = CACHE_DIR):
        self.cache_dir = cache_dir
        self._lock = threading.Lock()

    def _get_path(self, key: str) -> str:
        """Генерирует путь к файлу кеша на основе ключа."""
        hashed_key = hashlib.md5(key.encode('utf-8')).hexdigest()
        return os.path.join(self.cache_dir, f"{hashed_key}.json")

    def get(self, key: str) -> Optional[Any]:
        """Получает данные из кеша, если они свежие."""
        path = self._get_path(key)
        with self._lock:
            if not os.path.exists(path):
                return None
            
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                
                timestamp = data.get('timestamp', 0)
                if (time.time() - timestamp) > (CACHE_TTL_HOURS * 3600):
                    # Кеш протух
                    os.remove(path)
                    return None
                
                logger.debug(f"Cache HIT: {key[:30]}...")
                return data.get('payload')
            except Exception as e:
                logger.warning(f"Ошибка чтения кеша {path}: {e}")
                return None

    def set(self, key: str, payload: Any):
        """Сохраняет данные в кеш."""
        path = self._get_path(key)
        data = {
            "timestamp": time.time(),
            "key_raw": key,
            "payload": payload
        }
        with self._lock:
            try:
                with open(path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                logger.debug(f"Cache SET: {key[:30]}...")
            except Exception as e:
                logger.error(f"Не удалось записать кеш {path}: {e}")

class RateLimiter:
    """
    Контролирует частоту запросов к внешним ресурсам.
    Предотвращает блокировку по IP.
    """
    def __init__(self, calls_per_minute: int = 20):
        self.interval = 60.0 / calls_per_minute
        self.last_call = 0
        self._lock = threading.Lock()

    def wait(self):
        with self._lock:
            now = time.time()
            elapsed = now - self.last_call
            if elapsed < self.interval:
                sleep_time = self.interval - elapsed
                logger.debug(f"RateLimiter: Сон {sleep_time:.2f} сек.")
                time.sleep(sleep_time)
            self.last_call = time.time()

# Глобальные инстансы утилит
SYS_CACHE = CacheManager()
WEB_LIMITER = RateLimiter(calls_per_minute=15) # Осторожный серфинг

# ==============================================================================
# 10. КОНСТАНТЫ И МОДЕЛИ ДАННЫХ (Для WebSurfer)
# ==============================================================================

@dataclass
class SearchResultItem:
    title: str
    url: str
    snippet: str
    source_type: str = "web"

# ==============================================================================
# 11. ИНСТРУМЕНТЫ ПОИСКА И ИЗВЛЕЧЕНИЯ ДАННЫХ 
# ==============================================================================

class WebSurfer:
    """
    ENGINEERING WEBSURFER v10.0 (Stealth Mode).
    Использует API-обертку и эмуляцию человека для обхода блокировок.
    """
    def __init__(self): 
        self.session = requests.Session()

        # Улучшенная настройка ретраев: теперь реагируем на блокировку по частоте запросов (429)
        # и увеличили backoff_factor для более "вежливых" повторных попыток.
        retries = urllib3.util.retry.Retry(
            total=3, 
            backoff_factor=6,  # Паузы будут дольше (2с, 4с, 8с)
            status_forcelist=[429, 500, 502, 503, 504] # Добавлен код 429
        )
        adapter = requests.adapters.HTTPAdapter(max_retries=retries)
        self.session.mount('https://', adapter)
        self.session.mount('http://', adapter)
        
        # Защита от падения при старте, если API вики недоступен
        try:
            wikipedia.set_lang("ru")
        except Exception as e:
            logger.warning(f"Не удалось установить язык для Wikipedia: {e}")


    def _get_stealth_headers(self) -> Dict[str, str]:
        """
        Генерирует заголовки, имитирующие реальный браузер, из глобальной конфигурации.
        """
        #  списки берутся из глобального конфига CONFIG
        return {
            "User-Agent": random.choice(CONFIG.USER_AGENTS),
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.7",
            "Accept-Language": "ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7",
            "Accept-Encoding": "gzip, deflate, br",
            "Referer": random.choice(CONFIG.REFERERS),
            "Sec-Ch-Ua": '"Google Chrome";v="123", "Not:A-Brand";v="8", "Chromium";v="123"',
            "Sec-Ch-Ua-Mobile": "?0",
            "Sec-Ch-Ua-Platform": '"Windows"',
            "Sec-Fetch-Dest": "document",
            "Sec-Fetch-Mode": "navigate",
            "Sec-Fetch-Site": "cross-site",
            "Sec-Fetch-User": "?1",
            "Upgrade-Insecure-Requests": "1",
            "Connection": "keep-alive",
        }

    def search_google_simulated(self, query: str, limit: int = 5) -> List[SearchResultItem]:
        """
        Максимально скрытный поиск.
        """
        results = []
        
        # === МЕХАНИЗМ ЗАДЕРЖКИ (Anti-Ban) ===
        # Бот не должен делать запросы мгновенно. Спим от 2 до 5 секунд.
        sleep_time = random.uniform(2.5, 5.0)
        logger.info(f"🕵️ Stealth: Жду {sleep_time:.1f}с перед поиском...")
        time.sleep(sleep_time)

        # ПОПЫТКА 1: Библиотека DDGS (Самый надежный метод)
        try:
            with DDGS() as ddgs:
                ddg_gen = ddgs.text(
                    query, 
                    region="ru-ru", 
                    safesearch="off", 
                    timelimit="y",
                    max_results=limit + 2
                )
                
                for r in ddg_gen:
                    if len(results) >= limit: break
                    
                    link = r.get('href', '')
                    
                    title_raw = r.get('title', '')
                    # 1. Убираем переносы строк и лишние пробелы
                    title_clean = title_raw.replace('\n', ' ').strip()
                    # 2. Обрезаем слишком длинные заголовки, которые могут быть мусором
                    title = title_clean[:200] + '...' if len(title_clean) > 200 else title_clean
                    
                    body = r.get('body', '')

                    if not link or any(x in link.lower() for x in ['.pdf', '.doc', 'youtube.com', 'facebook.com']):
                        continue
                        
                    results.append(SearchResultItem(title=title, url=link, snippet=body))

            if results:
                logger.info(f"🔎 Успешный поиск (DDGS): {len(results)} результатов.")
                return results
        except Exception as e:
            logger.warning(f"⚠️ DDGS Blocked/Error: {e}")
            # Если словили блок, спим дольше перед Википедией
            time.sleep(3)

        # ПОПЫТКА 2: Википедия (Если DDG забанил)
        logger.warning(f"🛡️ Режим бункера: Переключаюсь на Wikipedia для '{query}'")
        try:
            wikipedia.set_lang("ru")
            # Используем try/except внутри цикла, чтобы одна ошибка не ломала всё
            search_res = wikipedia.search(query, results=3)
            
            for page_title in search_res:
                try:
                    page = wikipedia.page(page_title, auto_suggest=False)
                    results.append(SearchResultItem(
                        title=f"[Wiki] {page.title}",
                        url=page.url,
                        snippet=page.summary[:500],
                        source_type="encyclopedia"
                    ))
                except (wikipedia.DisambiguationError, wikipedia.PageError):
                    continue
            
            return results
        except Exception as e:
            logger.error(f"❌ Полный провал поиска: {e}")
            return []

    def get_page_content(self, url: str) -> str:
        """
        Умное скачивание страницы v3.0 с каскадным извлечением контента.
        """
        if "wikipedia.org" in url:
            try:
                title = url.split("/")[-1].replace('_', ' ')
                return wikipedia.page(title, auto_suggest=False).content[:25000]
            except Exception: return ""

        try:
            time.sleep(random.uniform(1.5, 3.5)) # Более "человеческая" пауза
            resp = self.session.get(url, headers=self._get_stealth_headers(), timeout=15, verify=False)
            
            if resp.status_code != 200: 
                logger.warning(f"Download failed with status {resp.status_code} for {url}")
                return ""
            
            # Улучшение: Авто-определение кодировки для борьбы с "кракозябрами"
            resp.encoding = resp.apparent_encoding
            
            soup = BeautifulSoup(resp.text, 'html.parser')
            
            # Улучшение: Более агрессивная очистка мусорных тегов
            for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'form', 'iframe', 'link', 'meta', 'noscript']):
                tag.decompose()
            
            # Улучшение: Каскадный поиск основного контента
            main_content_selectors = ['article', 'main', 'div#content', 'div.content', 'div.post-content', 'div.entry-content']
            main_content = None
            for selector in main_content_selectors:
                main_content = soup.select_one(selector)
                if main_content:
                    break

            if main_content:
                text = main_content.get_text(separator='\n', strip=True)
            elif soup.body:

                text = soup.body.get_text(separator='\n', strip=True)
            else:
                text = soup.get_text(separator='\n', strip=True)
            
            return self._sanitize_text(text)
            
        except Exception as e:
            logger.warning(f"Download Error for {url}: {e}")
            return ""

    def _sanitize_text(self, text: str) -> str:
        """Очистка текста v2.0: удаляет не только пробелы, но и семантический мусор."""
        
        # 1. Сначала базово чистим пробелы
        text = re.sub(r'[ \t]+', ' ', text)
        lines = text.split('\n')
        
        cleaned_lines = []
        # Паттерны для мусорных строк, которые нужно удалить
        junk_patterns = re.compile(
            r'^(?:Читать далее|Подписаться|Поделиться в|Источник:|Фото:|Автор:|Смотрите также|Похожие материалы)', 
            re.IGNORECASE
        )

        for line in lines:
            line = line.strip()
            # 2. Пропускаем пустые строки и слишком короткие 
            if not line or len(line) < 25:
                continue
            # 3. Пропускаем строки, которые являются явным мусором
            if junk_patterns.search(line):
                continue
            cleaned_lines.append(line)
        
        # Соединяем строки двойным переносом для лучшей читаемости LLM
        final_text = "\n\n".join(cleaned_lines)
        
        return final_text.strip()[:25000] # Ограничение на 25к символов 

class AcademicScholar:
    """
    Инструмент для работы с ArXiv.
    """
    def __init__(self):
        self.web = WebSurfer()

    def search_arxiv(self, query: str, limit: int = 5) -> List[SearchResultItem]:
        cache_key = f"search:arxiv:{query}"
        cached = SYS_CACHE.get(cache_key)
        if cached: return [SearchResultItem(**item) for item in cached]

        # Arxiv плохо ищет на русском, пропускаем
        if re.search('[а-яА-Я]', query): return []

        logger.info(f"📚 ArXiv Search: {query}")
        results = []
        try:
            client = arxiv.Client(page_size=limit, delay_seconds=3, num_retries=3)
            search = arxiv.Search(query=query, max_results=limit, sort_by=arxiv.SortCriterion.Relevance)
            
            for r in client.results(search):
                item = SearchResultItem(
                    title=r.title,
                    url=r.pdf_url,
                    snippet=r.summary.replace('\n', ' ')[:500],
                    source_type="arxiv",
                    published_date=str(r.published.date()),
                    authors=[a.name for a in r.authors]
                )
                results.append(item)
            
            if results: SYS_CACHE.set(cache_key, [r.to_dict() for r in results])
        except Exception: pass
        return results

    def download_pdf(self, url: str) -> str:
        # Используем метод веб-серфера для скачивания
        return self.web._process_pdf(url)

# ==============================================================================
# 12. ОБРАБОТКА ДОКУМЕНТОВ (The Hands)
# ==============================================================================

class DocumentForge:
    """
    ENGINEERING DOCUMENT FORGE v7.0 (Ultimate Hybrid).
    Класс-фабрика для создания профессионально отформатированных DOCX-документов.
    Объединяет надежность v5.0 с новыми возможностями v6.0 (LaTeX, Images, Smart Tables).
    Все методы статичны, так как класс не хранит состояние, а только предоставляет утилиты.
    """
    @staticmethod
    def create_styled_docx(master_plan, full_text_parts, bibliography, output_path: str) -> bool:
        """
        Создает финальный DOCX-документ с профессиональным форматированием.

        Args:
            master_plan (object): Объект, содержащий метаданные, такие как main_title и subtitle.
            full_text_parts (list[str] or str): Основной текст документа в виде списка строк или единой строки с разметкой Markdown.
            bibliography (list[dict]): Список библиографических ссылок, каждая в виде словаря с ключами 'title' и 'url'.
            output_path (str): Путь для сохранения итогового DOCX-файла.

        Returns:
            bool: True, если документ успешно создан и сохранен, иначе False.
        """
        # Создаем пустой документ Word.
        doc = docx.Document()

        # === 1. ЦЕНТРАЛИЗОВАННАЯ НАСТРОЙКА СТИЛЕЙ ===
        # Вызываем один метод, который определяет и настраивает все стили в документе.
        # Это обеспечивает консистентность и упрощает управление внешним видом.
        DocumentForge._define_document_styles(doc)

        # === 2. ФОРМИРОВАНИЕ ТИТУЛЬНОГО ЛИСТА ===
        # Используем отступы параграфов вместо пустых строк для точного контроля верстки.
        p_title = doc.add_paragraph()
        p_title.paragraph_format.first_line_indent = Inches(0) 
        p_title.paragraph_format.space_before = Pt(140)
        
        run_title = p_title.add_run(master_plan.main_title.upper())
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Настраиваем шрифт заголовка.
        font_title = run_title.font
        font_title.name = 'Times New Roman'
        font_title.bold = True
        font_title.size = Pt(16)

        # Добавляем подзаголовок, если он есть.
        if master_plan.subtitle:
            p_sub = doc.add_paragraph(master_plan.subtitle)
            p_sub.paragraph_format.first_line_indent = Inches(0)
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.runs[0].font.italic = True
            p_sub.runs[0].font.size = Pt(14)

        # Добавляем блок с датой внизу страницы.
        p_date = doc.add_paragraph()
        p_date.paragraph_format.first_line_indent = Inches(0)
        p_date.paragraph_format.space_before = Pt(250)  # Большой отступ сверху, чтобы сдвинуть текст вниз.
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_str = datetime.now().strftime("%d.%m.%Y")
        p_date.add_run(f"Аналитический отчет\n{date_str}")

        # Завершаем титульный лист разрывом страницы.
        doc.add_page_break()

        # === 3. РЕНДЕРИНГ ОСНОВНОГО КОНТЕНТА ===
        # Объединяем части текста в единую строку, если они переданы списком.
        full_text = "\n".join(full_text_parts) if isinstance(full_text_parts, list) else full_text_parts
        # Передаем весь текст в наш Markdown-парсер.
        DocumentForge._parse_markdown_to_docx(doc, full_text)

        # === 4. ФОРМИРОВАНИЕ СПИСКА ЛИТЕРАТУРЫ ===
        if bibliography:
            doc.add_page_break()
            doc.add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)

            # Умная дедупликация источников по URL, чтобы избежать повторов.
            seen_urls = set()
            unique_refs = []
            for ref in bibliography:
                url = ref.get('url', '').strip()
                # Добавляем источник, только если его URL уникален и не пуст.
                if url and url not in seen_urls:
                    seen_urls.add(url)
                    unique_refs.append(ref)

            current_date_str = datetime.now().strftime("%d.%m.%Y")
            for i, ref in enumerate(unique_refs, 1):
                title = ref.get('title', 'Электронный ресурс').replace('\n', ' ').strip()
                url = ref.get('url', '')
                # Используем специальный стиль 'Bibliography' для форматирования.
                p = doc.add_paragraph(style='Bibliography')
                p.add_run(f"{i}. {title} // URL: {url} (дата обращения: {current_date_str}).")

        # === 5. СОХРАНЕНИЕ ФАЙЛА ===
        try:
            doc.save(output_path)
            logger.info(f"Финальный документ успешно сохранен: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Критическая ошибка при сохранении DOCX-файла: {e}", exc_info=True)
            return False

    @staticmethod
    def _define_document_styles(doc):
        """
        Единый центр управления стилями документа.
        Исправлено: безопасные гиперссылки и черные подписи к рисункам.
        """
        # --- 1. Базовый стиль 'Normal' ---
        style_normal = doc.styles['Normal']
        font = style_normal.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        font.color.rgb = RGBColor(0, 0, 0)

        p_format = style_normal.paragraph_format
        p_format.line_spacing = 1.15
        p_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_format.first_line_indent = Inches(0.5)
        p_format.space_after = Pt(6)

        # --- 2. Настройка встроенных стилей Word ---
        # Заголовки
        for level in range(1, 5):
            style_heading = doc.styles[f'Heading {level}']
            style_heading.base_style = style_normal
            font = style_heading.font
            font.name = 'Times New Roman'
            font.color.rgb = RGBColor(0, 0, 0)
            font.bold = True
            font.size = Pt(16 - level * 1.5)
            p_format = style_heading.paragraph_format
            p_format.first_line_indent = Inches(0)
            p_format.space_before = Pt(18)
            p_format.space_after = Pt(6)

        # Гиперссылки (FIXED)
        try:
            style_hyperlink = doc.styles['Hyperlink']
        except KeyError:
            style_hyperlink = doc.styles.add_style('Hyperlink', WD_STYLE_TYPE.CHARACTER)
        
        font = style_hyperlink.font
        font.color.rgb = RGBColor(0, 0, 255)
        font.underline = True

        # --- 3. Создание кастомных стилей ---
        
        if 'CodeBlock' not in doc.styles:
            style = doc.styles.add_style('CodeBlock', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_normal
            style.font.name = 'Consolas'
            style.font.size = Pt(10)
            p_format = style.paragraph_format
            p_format.first_line_indent = Inches(0)
            p_format.left_indent = Inches(0.5)
            p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if 'BlockQuote' not in doc.styles:
            style = doc.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_normal
            style.font.italic = True
            p_format = style.paragraph_format
            p_format.left_indent = Inches(0.5)
            p_format.first_line_indent = Inches(0)

        # Caption (FIXED: Принудительно черный цвет)
        if 'Caption' not in doc.styles:
            style = doc.styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
        else:
            style = doc.styles['Caption']
            
        style.base_style = style_normal
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        style.font.italic = True
        style.font.color.rgb = RGBColor(0, 0, 0) # Черный цвет
        
        p_format = style.paragraph_format
        p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_format.first_line_indent = Inches(0)
        p_format.space_before = Pt(6)

        # Списки
        if 'List Bullet' not in doc.styles:
            style = doc.styles.add_style('List Bullet', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_normal
            p_format = style.paragraph_format
            p_format.first_line_indent = Inches(0)
            p_format.left_indent = Inches(0.5)

        if 'List Number' not in doc.styles:
            style = doc.styles.add_style('List Number', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_normal
            p_format = style.paragraph_format
            p_format.first_line_indent = Inches(0)
            p_format.left_indent = Inches(0.5)
            
        if 'Bibliography' not in doc.styles:
            style = doc.styles.add_style('Bibliography', WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = style_normal
            style.font.size = Pt(11)
            p_format = style.paragraph_format
            p_format.first_line_indent = Inches(0)
            p_format.hanging_indent = Inches(0.5)
            p_format.space_after = Pt(4)

    @staticmethod
    def _format_latex_as_text(latex_str: str) -> str:
        """
        ULTIMATE EDITION: Конвертер LaTeX в текст.
        Гарантирует отсутствие программного кода в DOCX.
        Использует стратегию "послойного снятия" разметки.
        """
        if not latex_str:
            return ""

        s = latex_str.strip()

        # 1. Удаляем внешние границы формул
        # Удаляем \[ ... \], $$ ... $$, \( ... \)
        s = re.sub(r'^\\\[|\\\]$', '', s)
        s = re.sub(r'^\$\$|\$\$$', '', s)
        s = re.sub(r'^\\\(|\\\)$', '', s)

        # 2. Переводим спецсимволы в Unicode (пока структура цела)
        replacements = {
            r'\\approx': '≈', r'\\ne': '≠', r'\\neq': '≠',
            r'\\le': '≤', r'\\leq': '≤', r'\\ge': '≥', r'\\geq': '≥',
            r'\\times': '×', r'\\cdot': '·', r'\\pm': '±', 
            r'\\infty': '∞', r'\\rightarrow': '→', r'\\leftarrow': '←',
            r'\\%': '%', r'\\EUR': '€', r'\\USD': '$', r'\\deg': '°',
            r'\\text\s*\{': '', # Удаляем начало команды \text{
        }
        for pattern, char in replacements.items():
            s = re.sub(pattern, char, s, flags=re.IGNORECASE)

        # 3. УНИЧТОЖЕНИЕ КОМАНД ОФОРМЛЕНИЯ
        # Удаляем сами команды, но оставляем содержимое скобок.
        # \textbf{Word} -> {Word}
        commands_to_kill = [
            r'\\text', r'\\textbf', r'\\textit', r'\\mathrm', 
            r'\\mathbf', r'\\mbox', r'\\quad', r'\\,'
        ]
        for cmd in commands_to_kill:
            s = re.sub(cmd, '', s)

        # 4. ОБРАБОТКА ДРОБЕЙ (Циклическая)
        # \frac{A}{B} -> (A) / (B)
        # Запускаем цикл 5 раз, чтобы раскрыть вложенные дроби: \frac{\frac{a}{b}}{c}
        for _ in range(5):
            # Ищем паттерн \frac{...}{...}. 
            # Используем [^{}] чтобы найти ближайшую закрывающую скобку
            new_s = re.sub(r'\\frac\s*\{([^{}]+)\}\s*\{([^{}]+)\}', r'(\1) / (\2)', s)
            
            if new_s == s: 
                # План Б: Если остались \frac, но скобки сложные, просто удаляем \frac
                s = s.replace(r'\frac', '') 
                break
            s = new_s

        # 5. Индексы и степени
        s = re.sub(r'_\{([^}]+)\}', r'_\1', s) # x_{i} -> x_i
        s = re.sub(r'\^\{([^}]+)\}', r'^\1', s) # x^{2} -> x^2

        # 6. ФИНАЛЬНАЯ ЗАЧИСТКА (FAIL-SAFE)
        # Это гарантия того, что пользователь не увидит код.
        # Если после всех манипуляций остались символы синтаксиса — удаляем их.
        
        # Удаляем все обратные слеши
        s = s.replace('\\', '')
        
        # Удаляем все фигурные скобки
        s = s.replace('{', '').replace('}', '')
        
        # Удаляем квадратные скобки, если они остались от [Code]
        s = s.replace('[', '').replace(']', '')

        # 7. Косметика текста
        # Убираем двойные пробелы
        s = re.sub(r'\s+', ' ', s).strip()
        # Убираем двойные скобки: ((A)) -> (A), возникшие при раскрытии дробей
        s = s.replace('((', '(').replace('))', ')')
        # Исправляем пробелы вокруг слеша: ( A ) / ( B ) -> (A) / (B)
        s = re.sub(r'\(\s+', '(', s)
        s = re.sub(r'\s+\)', ')', s)

        return s

    @staticmethod
    def _parse_markdown_to_docx(doc, text: str):
        """
        Машина состояний для парсинга текста с Markdown-разметкой в DOCX.
        Обрабатывает текст построчно, переключая состояния (NORMAL, CODE, TABLE).
        """
        lines = text.split('\n')
        state = "NORMAL"  # Текущее состояние парсера
        buffer = []       # Буфер для многострочных элементов (код, таблицы)

        # Компилируем регулярные выражения один раз для повышения производительности
        latex_block_pattern = re.compile(r'^\s*(\$\$|\\\[)(.*?)(\$\$|\\\])\s*$', re.DOTALL)
        img_pattern = re.compile(r'^\s*!\[(.*?)\]\((.*?)\)\s*$')
        num_list_pattern = re.compile(r'^\s*\d+\.\s+')
        # Паттерн для линий-разделителей (---, ***, ___)
        separator_pattern = re.compile(r'^[-*_]{3,}\s*$')

        for line in lines:
            stripped = line.strip()

            # --- 1. УПРАВЛЕНИЕ СОСТОЯНИЯМИ МНОГОСТРОЧНЫХ БЛОКОВ ---
            
            # Блок кода (```)
            if stripped.startswith("```"):
                if state == "NORMAL":
                    state = "CODE"
                    continue  # Начало блока, пропускаем строку ```
                elif state == "CODE":
                    # Конец блока: рендерим содержимое буфера
                    code_text = '\n'.join(p.rstrip('\n\r') for p in buffer)
                    if code_text:
                        doc.add_paragraph(code_text, style='CodeBlock')
                    buffer = []
                    state = "NORMAL"
                    continue # Пропускаем строку ```

            if state == "CODE":
                buffer.append(line) # Накапливаем строки кода в буфер
                continue

            # Таблица (|...|)
            is_table_line = stripped.startswith("|") and stripped.count("|") > 1
            
            if state == "TABLE" and not is_table_line:
                # Если мы были в таблице, а текущая строка - нет, значит таблица закончилась.
                DocumentForge._render_table(doc, buffer)
                buffer = []
                state = "NORMAL"
                # ВАЖНО: Не используем 'continue', т.к. текущую строку нужно обработать дальше.

            if is_table_line:
                if '---' in stripped: # Пропускаем разделитель Markdown-таблиц
                    continue
                if state != "TABLE":
                    state = "TABLE" # Входим в состояние "TABLE"
                buffer.append(stripped)
                continue # Строка обработана, переходим к следующей

            # --- 2. ОБРАБОТКА ОДНОСТРОЧНЫХ ЭЛЕМЕНТОВ (когда state == "NORMAL") ---

            if not stripped:
                continue
                
            # [ИСПРАВЛЕНИЕ] Игнорируем разделители (---)
            if separator_pattern.match(stripped):
                continue
            
            # Заголовки (#, ##, ...)
            if stripped.startswith('#'):
                level = len(stripped) - len(stripped.lstrip('#'))
                clean_text = stripped.lstrip('#').strip()
                # Удаляем ручную нумерацию (1. Введение -> Введение), доверяя стилям или чистя мусор
                clean_text = re.sub(r'^\d+(\.\d+)*\.?\s*', '', clean_text)
                doc.add_heading(clean_text, level=min(level, 4)) 
                continue

            # === ФИКС: Псевдо-заголовки (**Текст**) ===
            # Если строка целиком обернута в **, считаем её подзаголовком
            if stripped.startswith('**') and stripped.endswith('**') and len(stripped) > 4:
                clean_text = stripped[2:-2].strip()
                # Создаем параграф, делаем его жирным и добавляем отступ сверху
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12) # Отступ как у заголовка
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.keep_with_next = True # Не отрывать от следующего текста
                
                run = p.add_run(clean_text)
                run.font.bold = True
                continue

            # Цитаты (>)
            if stripped.startswith('>'):
                clean_text = stripped.lstrip('>').strip()
                p = doc.add_paragraph(style='BlockQuote')
                DocumentForge._add_rich_text(p, clean_text)
                continue
            
            # Изображения (![alt](path))
            img_match = img_pattern.match(stripped)
            if img_match:
                alt_text, img_path = img_match.groups()
                # Убираем кавычки из пути, если есть
                img_path = img_path.strip('"\'') 
                
                if os.path.exists(img_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.paragraph_format.first_line_indent = Inches(0)
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(img_path, width=Inches(6.0))
                        
                        if alt_text: 
                            # === ФИКС ДУБЛИРОВАНИЯ ===
                            # Проверяем, есть ли уже слово "Рисунок" в начале
                            if re.match(r'^(Рисунок|Figure|Рис\.)', alt_text, re.IGNORECASE):
                                caption_text = alt_text
                            else:
                                caption_text = f"Рисунок: {alt_text}"
                            
                            doc.add_paragraph(caption_text, style='Caption')
                            # =========================
                    except Exception as e:
                        doc.add_paragraph(f"[Ошибка загрузки изображения: {e}]", style='CodeBlock')
                else:
                    doc.add_paragraph(f"[Изображение не найдено: {img_path}]", style='CodeBlock')
                continue

            # Блоки LaTeX ($$ ... $$)
            # [FIX] Если строка просто содержит $$, считаем её формулой, даже если регулярка не совпала идеально
            if '$$' in stripped or latex_block_pattern.match(line):
                # Убираем $$ и отправляем в чистильщик
                raw_formula = line.replace('$$', '').strip()
                formatted_formula = DocumentForge._format_latex_as_text(raw_formula)
                
                # Добавляем в документ
                p = doc.add_paragraph(style='CodeBlock') # Или Normal, если хотите
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Важно: используем обычный шрифт, не курсив, чтобы цифры были четкими
                run = p.add_run(formatted_formula)
                run.font.italic = True
                continue

            # Маркированные списки (* или -)
            if stripped.startswith(('- ', '* ')):
                clean_text = stripped[2:]
                p = doc.add_paragraph(style='List Bullet')
                DocumentForge._add_rich_text(p, clean_text)
                continue

            # Нумерованные списки (1., 2., ...)
            if num_list_pattern.match(stripped):
                clean_text = num_list_pattern.sub('', stripped, count=1)
                p = doc.add_paragraph(style='List Number')
                DocumentForge._add_rich_text(p, clean_text)
                continue
            
            # Обычный параграф
            # [ИСПРАВЛЕНИЕ] Более агрессивная очистка от LaTeX мусора
            clean_line = stripped.replace(r'\[', '').replace(r'\]', '').replace(r'\(', '').replace(r'\)', '')
            p = doc.add_paragraph() # Используем стиль 'Normal' по умолчанию
            DocumentForge._add_rich_text(p, clean_line)

        # --- 3. ЗАВЕРШЕНИЕ РАБОТЫ ---
        if buffer:
            if state == "TABLE":
                DocumentForge._render_table(doc, buffer)
            elif state == "CODE":
                code_text = '\n'.join(p.rstrip('\n\r') for p in buffer)
                if code_text:
                    doc.add_paragraph(code_text, style='CodeBlock')

    @staticmethod
    def _render_table(doc, rows: list[str]):
        """
        Рендерит Markdown таблицу в стилизованную таблицу DOCX.
        """
        # 1. Парсинг и очистка данных
        # Пример: '| a | b |' -> ['', ' a ', ' b ', ''] -> ['a', 'b']
        parsed_rows = [
            [cell.strip() for cell in row.split('|')][1:-1]
            for row in rows if row.strip().startswith('|')
        ]
        # Убираем пустые строки (могли остаться от разделителя |---|)
        parsed_rows = [row for row in parsed_rows if any(cell for cell in row)]
        
        if not parsed_rows: return

        # 2. Создание и настройка таблицы
        num_cols = max(len(row) for row in parsed_rows)
        if num_cols == 0: return

        table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
        
        # Безопасное присвоение стиля таблицы
        try:
            table.style = 'Table Grid'
        except KeyError:
            # Если стиля 'Table Grid' нет, пробуем стандартный или оставляем без стиля
            try:
                table.style = 'Normal Table'
            except KeyError:
                pass # Оставляем таблицу без явного стиля
                
        table.autofit = True
        # 3. Заполнение и стилизация ячеек
        for i, row_data in enumerate(parsed_rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j >= len(row_cells): continue
                
                cell = row_cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Выравниваем по центру для красоты
                cell.text = "" # Очищаем стандартный параграф Word
                
                # Добавляем свой параграф для полного контроля над форматированием
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                DocumentForge._add_rich_text(p, cell_text.strip())

                # Особое форматирование для заголовка (первая строка)
                if i == 0:
                    for run in p.runs: # Делаем весь текст в ячейке жирным
                        run.font.bold = True
                    
                    # Добавляем заливку фона (низкоуровневая магия oxml)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:fill'), 'F2F2F2') # Светло-серый цвет
                    tc_pr.append(shd)

    @staticmethod
    def _add_rich_text(paragraph, text: str):
        """
        Парсит инлайн-стили Markdown (**bold**, *italic*, `code`, [link](url))
        и добавляет их в параграф как форматированный текст.
        Исправлена ошибка с синим цветом и остающимися звездочками.
        """
        # Регулярное выражение. Важен порядок: сначала ищем двойные звездочки!
        # re.DOTALL позволяет . захватывать и переносы строк
        pattern = re.compile(
            r'(?P<bold>\*\*(?P<bold_text>.+?)\*\*)|'             # **Bold**
            r'(?P<italic>(?<!\*)\*(?P<italic_text>[^*]+?)\*(?!\*))|' # *Italic* (исключает **)
            r'(?P<code>`(?P<code_text>.+?)`)|'                   # `Code`
            r'(?P<link>\[(?P<link_text>[^\]]+)\]\((?P<url>[^\)]+)\))', # [Link](url)
            re.DOTALL
        )
        
        last_end = 0 
        
        # Итерация по всем найденным совпадениям
        for match in pattern.finditer(text):
            start, end = match.span()
            
            # 1. Добавляем обычный текст ДО совпадения (если есть)
            if start > last_end:
                run = paragraph.add_run(text[last_end:start])
                # Явно задаем черный цвет для обычного текста, чтобы избежать "синевы"
                run.font.color.rgb = RGBColor(0, 0, 0)

            group_name = match.lastgroup
            
            # 2. Обработка стилей
            if group_name == 'bold':
                # Берем текст ИЗНУТРИ звездочек
                clean_text = match.group('bold_text')
                run = paragraph.add_run(clean_text)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Принудительно черный
            
            elif group_name == 'italic':
                clean_text = match.group('italic_text')
                run = paragraph.add_run(clean_text)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0) # Принудительно черный
                
            elif group_name == 'code':
                clean_text = match.group('code_text')
                run = paragraph.add_run(clean_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(10)
                # Фон для кода
                rpr = run._r.get_or_add_rPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:fill'), 'F1F1F1') # Светло-серый фон
                rpr.append(shd)
                run.font.color.rgb = RGBColor(0, 0, 0)

            elif group_name == 'link':
                link_text = match.group('link_text')
                url = match.group('url')
                
                # Создание кликабельной ссылки через Oxml
                try:
                    part = paragraph.part
                    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

                    hyperlink = OxmlElement('w:hyperlink')
                    hyperlink.set(qn('r:id'), r_id)

                    run_element = OxmlElement('w:r')
                    rpr = OxmlElement('w:rPr')
                    
                    # Применяем стиль, который мы безопасно создали в _define_document_styles
                    style_element = OxmlElement('w:rStyle')
                    style_element.set(qn('w:val'), 'Hyperlink')
                    rpr.append(style_element)
                    
                    text_element = OxmlElement('w:t')
                    text_element.text = link_text
                    
                    run_element.append(rpr)
                    run_element.append(text_element)
                    hyperlink.append(run_element)
                    paragraph._p.append(hyperlink)
                except Exception:
                    # Если ссылка сломалась, выводим просто текст, чтобы не терять контент
                    run = paragraph.add_run(link_text)
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.font.underline = True

            last_end = end
            
        # 3. Добавляем "хвост" текста после последнего совпадения
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.color.rgb = RGBColor(0, 0, 0)

# ==============================================================================
# 13. БАЗА ЗНАНИЙ (ГРАФ + ВЕКТОРЫ)
# ==============================================================================

@dataclass
class KnowledgeNode:
    id: str
    type: str # 'concept', 'fact', 'source'
    content: str
    meta: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None

class KnowledgeGraphEngine:
    """
    Графовая память системы. Хранит связи между фактами.
    Использует NetworkX для графа и JSON для персистентности.
    """
    def __init__(self, db_path: str = CONFIG.ARTIFACTS_DIR):
        self.graph = nx.DiGraph()
        self.db_path = os.path.join(db_path, "knowledge_graph.json")
        self._load() 
        self.lock = threading.Lock()

    def add_node(self, content: str, node_type: str, meta: Dict = None) -> str:
        """Добавляет узел. Возвращает ID."""
        node_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, content))[:12]
        
        with self.lock:
            if node_id not in self.graph:
                self.graph.add_node(
                    node_id, 
                    type=node_type, 
                    content=content, 
                    meta=meta or {},
                    created_at=time.time()
                )
                logger.debug(f"KnowledgeGraph: Добавлен узел [{node_type}] {node_id}")
        return node_id

    def add_edge(self, source_id: str, target_id: str, relation: str):
        """Добавляет связь."""
        with self.lock:
            if self.graph.has_node(source_id) and self.graph.has_node(target_id):
                self.graph.add_edge(source_id, target_id, relation=relation)

    def find_related(self, node_id: str, depth: int = 1) -> List[Dict]:
        """Находит связанные концепции."""
        if node_id not in self.graph:
            return []
        
        # BFS поиск соседей
        related = []
        try:
            neighbors = list(nx.bfs_tree(self.graph, node_id, depth_limit=depth))
            for n_id in neighbors:
                if n_id == node_id: continue
                node_data = self.graph.nodes[n_id]
                related.append({
                    "id": n_id,
                    "content": node_data['content'],
                    "type": node_data['type']
                })
        except Exception as e:
            logger.error(f"Ошибка поиска в графе: {e}")
        
        return related

    def get_all_text_context(self) -> str:
        """Выгружает все знания в один текст (для контекста LLM)."""
        lines = []
        for n, d in self.graph.nodes(data=True):
            lines.append(f"[{d.get('type', 'info').upper()}] {d.get('content', '')}")
        return "\n".join(lines)

    def save(self):
        """Сохраняет граф на диск."""
        with self.lock:
            data = nx.node_link_data(self.graph)
            try:
                with open(self.db_path, 'w', encoding='utf-8') as f:
                    json.dump(data, f, ensure_ascii=False, indent=2)
                logger.info(f"Граф знаний сохранен. Узлов: {self.graph.number_of_nodes()}")
            except Exception as e:
                logger.error(f"Ошибка сохранения графа: {e}")

    def _load(self):
        """Загружает граф с диска."""
        if os.path.exists(self.db_path):
            try:
                with open(self.db_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                self.graph = nx.node_link_graph(data)
                logger.info(f"Граф знаний загружен. Узлов: {self.graph.number_of_nodes()}")
            except Exception as e:
                logger.error(f"Ошибка загрузки графа, создан новый: {e}")

    def search_text(self, query: str) -> List[Dict]:
        """
        Умный полнотекстовый поиск по графу знаний.
        Использует эвристику пересечения слов (Bag of Words) для ранжирования.
        """
        if not query:
            return []

        results = []
        # Нормализация запроса: нижний регистр, разбиение на слова
        query_words = set(re.findall(r'\w+', query.lower()))
        # Фильтрация стоп-слов (коротких предлогов и союзов) для русского и английского
        filtered_query_words = {w for w in query_words if len(w) > 3}
        
        if not filtered_query_words:
            # Если остались только короткие слова, ищем как есть
            filtered_query_words = query_words

        for n, d in self.graph.nodes(data=True):
            content = d.get('content', '')
            if not content: continue
            
            content_lower = content.lower()
            score = 0
            matches = []
            
            # Считаем совпадения
            for word in filtered_query_words:
                if word in content_lower:
                    score += 1
                    matches.append(word)
            
            # Дополнительный вес за совпадение в заголовке/мете (если есть)
            meta_title = d.get('meta', {}).get('title', '').lower()
            if meta_title:
                for word in filtered_query_words:
                    if word in meta_title:
                        score += 2 # Заголовок важнее
            
            # Если есть хотя бы одно совпадение
            if score > 0:
                # Нормализуем скор по длине (чтобы короткие совпадения в заголовках ценились)
                # Но отдаем приоритет количеству уникальных совпавших слов
                results.append({
                    "id": n,
                    "content": content,
                    "meta": d.get('meta', {}),
                    "type": d.get('type', 'unknown'),
                    "score": score,
                    "matches": matches
                })
        
        # Сортировка: сначала по скору (кол-во совпадений), потом по типу (evidence важнее)
        results.sort(key=lambda x: (x['score'], 1 if x['type'] == 'evidence' else 0), reverse=True)
        
        # Возвращаем топ-7 самых релевантных
        return results[:7]


# ==============================================================================
# 14. ИНТЕГРАТОР ИНСТРУМЕНТОВ (Facede Pattern)
# ==============================================================================

class GenesisToolbox:
    """
    Единая точка входа для Агентов.
    Агенты не должны создавать экземпляры классов выше напрямую.
    """
    def __init__(self):
        self.web = WebSurfer()
        self.scholar = AcademicScholar()
        self.graph = KnowledgeGraphEngine()
        self.doc_forge = DocumentForge()
        self.start_time = time.time()

    def research_topic(self, topic: str) -> str:
        """
        Выполняет комплексное исследование:
        1. Поиск в ArXiv
        2. Поиск в Web
        3. Сохранение в Граф
        4. Возврат сводки
        """
        logger.info(f"TOOLBOX: Запуск исследования по теме '{topic}'")
        
        report = []
        
        # 1. Arxiv
        arxiv_results = self.scholar.search_arxiv(topic, limit=3)
        report.append(f"Найдено в ArXiv: {len(arxiv_results)} статей.")
        for item in arxiv_results:
            self.graph.add_node(
                f"Paper: {item.title}\nAbstract: {item.snippet}", 
                "paper", 
                {"url": item.url, "authors": item.authors, "title": item.title} # Added title for better search
            )
        
        # 2. Web
        web_results = self.web.search_google_simulated(topic, limit=3)
        report.append(f"Найдено в Web: {len(web_results)} страниц.")
        for item in web_results:
            self.graph.add_node(
                f"WebPage: {item.title}\nSnippet: {item.snippet}", 
                "web_snippet",
                {"url": item.url, "title": item.title} # Added title
            )
            
            # Глубокое чтение первого результата
            if item == web_results[0]:
                full_text = self.web.get_page_content(item.url)
                if full_text:
                    node_id = self.graph.add_node(full_text[:2000], "document", {"url": item.url, "title": item.title})
                    report.append(f"Прочитана страница: {item.url} (ID: {node_id})")

        self.graph.save()
        return "\n".join(report)

# ==============================================================================
# 15. БАЗОВЫЕ КЛАССЫ АГЕНТОВ 
# ==============================================================================

class AgentStatus(str, Enum):
    IDLE = "idle"
    THINKING = "thinking"
    WORKING = "working"
    WAITING = "waiting"
    ERROR = "error"
    DONE = "done"

@dataclass
class AgentThought:
    """Единица мыслительного процесса агента (для лога)."""
    step: str
    content: str
    timestamp: float = field(default_factory=time.time)

class BaseAgent:
    """
    Абстрактный класс для всех агентов.
    """
    def __init__(self, name: str, role_prompt: str, engine: LLMEngine, tools: GenesisToolbox):
        self.name = name
        self.role_prompt = role_prompt
        self.llm = engine
        self.tools = tools
        self.status = AgentStatus.IDLE
        self.thought_chain: List[AgentThought] = []
        
        self.logger = logging.getLogger(f"Agent.{name}")
        self.logger.setLevel(logging.INFO)

    def log_thought(self, step: str, content: str):
        """Записывает 'мысль' агента."""
        thought = AgentThought(step, content)
        self.thought_chain.append(thought)
        self.logger.info(f"[{step.upper()}] {content[:100]}...")


# ==============================================================================
# 16. АРХИТЕКТОР (THE ARCHITECT) - СТРАТЕГИЧЕСКОЕ ПЛАНИРОВАНИЕ 
# ==============================================================================

class ArchitectAgent(BaseAgent):
    """
    Отвечает за создание глубокого и детализированного мастер-плана за один проход.
    Версия 7.0 "Hybrid Counter".
    """
    def __init__(self, engine, tools):
        super().__init__("Architect", PromptsLibrary.ARCHITECT_ROLE, engine, tools)
        
        self.EVIDENCE_CHUNK_SIZE = 6000  # Размер куска текста для анализа фактов
        self.MAX_CONTEXT_CHARS = 3000    # Сколько символов из прошлого помнить
        self.TEMP_CREATIVE = 0.7         # Температура креативности
        self.MIN_TEXT_LENGTH = 100       # Минимальная длина главы

    def create_master_plan(self, topic: str) -> Optional[ArticleMasterPlan]:
        
        # 1. АНАЛИЗ ВВОДНЫХ ДАННЫХ (Определение режима работы и целевого числа глав)
        
        # Режим 1: Проверка на готовый, нумерованный план от пользователя
        # Ищем паттерны: "1.", "2.", "3. " в начале строк
        detected_chapters = re.findall(r'^\s*\d+\.\s+', topic, re.MULTILINE)
        
        if len(detected_chapters) >= 2:
            # Если нашли явный нумерованный список, это режим "ИЗВЛЕЧЕНИЕ"
            target_count = len(detected_chapters)
            mode = "EXTRACT"
            self.log_thought("ANALYSIS", f"Обнаружен готовый план в запросе ({target_count} пунктов). Режим: СТРОГОЕ ИЗВЛЕЧЕНИЕ.")
        else:
            # Режим 2: Используем значение из CONFIG (генерация с нуля)
            target_count = CONFIG.TARGET_CHAPTER_COUNT
            mode = "GENERATE"
            self.log_thought("ANALYSIS", f"Детальный план не обнаружен. Режим: ГЕНЕРАЦИЯ с нуля ({target_count} глав из CONFIG).")
            
        # 2. ПОДГОТОВКА БАЗОВОГО ПРОМПТА
        if mode == "EXTRACT":
            prompt = (
                f"ИСХОДНЫЙ ДЕТАЛЬНЫЙ ПЛАН СТАТЬИ:\n---\n{topic}\n---\n\n"
                f"ЗАДАЧА: Ты — системный анализатор. Твоя задача — не придумывать ничего нового, а ТОЧНО ИЗВЛЕЧЬ структуру из текста выше и преобразовать ее в JSON-формат.\n"
                f"КРИТИЧЕСКИ ВАЖНО: Исходный план содержит ровно {target_count} разделов. Твой JSON-ответ ОБЯЗАН содержать ровно {target_count} глав в списке 'chapters'.\n"
                f"Обязательно заполни все поля для каждой главы: `title`, `purpose`, `core_thesis`, `key_points`, `narrative_link`.\n"
                f"Если в исходном плане нет поля `purpose` или `narrative_link`, сгенерируй их на основе контекста главы.\n"
                f"Верни только JSON."
            )
        else:
            # Режим GENERATE (твой оригинальный код, но с усиленной инструкцией)
            prompt = (
                f"ТЕМА ИССЛЕДОВАНИЯ:\n---\n{topic}\n---\n\n"
                f"ЗАДАЧА: Создай исчерпывающий мастер-план для аналитической статьи/книги по этой теме.\n"
                f"КРИТИЧЕСКИ ВАЖНО: План должен состоять СТРОГО ИЗ {target_count} ГЛАВ. Не больше и не меньше. Убедись, что в итоговом JSON в списке 'chapters' находится ровно {target_count} объектов.\n\n"
                f"Для КАЖДОЙ из {target_count} глав заполни все поля: `title`, `purpose`, `core_thesis`, `key_points`, `narrative_link`.\n\n"
                f"Выполни всю работу за один раз. Верни только JSON."
            )
        
        # 3. ЦИКЛ ГЕНЕРАЦИИ И САМОКОРРЕКЦИИ (Гарантия качества)
        master_plan = None
        
        # Даем 3 попытки на извлечение/генерацию
        for attempt in range(3):
            if attempt > 0:
                # Если мы здесь, значит предыдущая попытка была неудачной
                prev_count = len(master_plan.chapters) if master_plan and master_plan.chapters else 0
                self.log_thought("RETRY_ARCHITECT", f"Попытка {attempt+1}: Модель создала {prev_count} глав вместо {target_count}. Требую исправить.")
                
                # Дополняем промпт требованием исправить ошибку
                correction_request = (
                    f"\n\n*** СИСТЕМНАЯ ОШИБКА ПРЕДЫДУЩЕЙ ПОПЫТКИ: "
                    f"Ты сгенерировал {prev_count} глав. "
                    f"А НУЖНО РОВНО {target_count}. ПЕРЕДЕЛАЙ ПЛАН! ***"
                )
                current_prompt = prompt + correction_request
            else:
                current_prompt = prompt

            # Вызываем LLM
            master_plan = self.llm.generate_structured(
                system_prompt=self.role_prompt,
                user_prompt=current_prompt,
                response_model=ArticleMasterPlan,
                complexity=ComplexityLevel.ANALYTICAL,
                temperature=0.7 + (attempt * 0.1) # Чуть повышаем температуру для вариативности
            )

            # Проверка результата
            if master_plan and master_plan.chapters and len(master_plan.chapters) == target_count:
                # Успех: найдено точное количество глав!
                break 

        # 4. ФИНАЛЬНЫЙ АУДИТ
        
        if not master_plan or not master_plan.chapters:
            self.log_thought("FAIL", "Модель не смогла сгенерировать или извлечь структуру. Возвращаю None.")
            return None

        # Логируем, если количество глав все равно не совпало (даже после 3 попыток)
        if len(master_plan.chapters) != target_count:
            self.log_thought("WARN", f"Не удалось добиться точного совпадения ({len(master_plan.chapters)}/{target_count}). Продолжаем с тем, что есть.")
        else:
            self.log_thought("SUCCESS", f"Мастер-план на {len(master_plan.chapters)} глав успешно создан/извлечен.")

        # Сохраняем в граф только финальный результат
        self.tools.graph.add_node(
            content=master_plan.model_dump_json(indent=2),
            node_type="master_plan",
            meta={"topic": topic, "title": master_plan.main_title}
        )
        self.tools.graph.save() 
        
        return master_plan

# ==============================================================================
# 17. ИССЛЕДОВАТЕЛЬ (THE RESEARCHER) - ПОИСК И ФИЛЬТРАЦИЯ 
# ==============================================================================

class ResearcherAgent(BaseAgent):
    """
    Отвечает за поиск фактов. Версия 7.0 "Evidence Hunter".
    Проводит целевой поиск доказательств для каждой главы из мастер-плана.
    """
    def __init__(self, engine, tools):
        # Используем EVIDENCE_DISTILLER_PROTOCOL как основной протокол агента для анализа источника
        super().__init__("Researcher", PromptsLibrary.EVIDENCE_DISTILLER_PROTOCOL, engine, tools)

    class SearchQueryPlan(BaseModel):
        """
        Модель для извлечения поисковых запросов.
        Pydantic v2 отлично справляется с парсингом даже "грязных" списков.
        """
        queries: List[str] = Field(default_factory=list, description="Список поисковых запросов")

        @field_validator('queries', mode='before')
        @classmethod
        def robust_query_extractor(cls, v: Any) -> List[str]:
            """
            Магический фильтр. Превращает любую структуру данных в список строк.
            Решает проблему: [{"type": "confirm", "query": "text"}] -> ["text"]
            """
            extracted = []
            
            # 1. Если пришел не список (например, один словарь), оборачиваем в список
            if not isinstance(v, list):
                v = [v]

            # 2. Рекурсивный/Итеративный разбор элементов
            for item in v:
                if isinstance(item, str):
                    # Если это просто строка — берем, но чистим от нумерации "1. "
                    clean = re.sub(r'^\d+[\).]\s*', '', item).strip()
                    if len(clean) > 5: extracted.append(clean)
                
                elif isinstance(item, dict):
                    # Эвристика: ищем поле, похожее на запрос
                    # Приоритет ключей, которые обычно используют модели
                    target_keys = ['query', 'q', 'text', 'content', 'value', 'search_term', 'topic']
                    
                    found_val = None
                    # Проход 1: Точное совпадение ключа
                    for key in target_keys:
                        if key in item and isinstance(item[key], str):
                            found_val = item[key]
                            break
                    
                    # Проход 2: Если ключи не найдены, ищем любую длинную строку в значениях
                    if not found_val:
                        for val in item.values():
                            if isinstance(val, str) and len(val) > 10 and '?' not in val: # Игнорируем вопросы-метаданные
                                found_val = val
                                break
                    
                    if found_val:
                        clean = re.sub(r'^\d+[\).]\s*', '', found_val).strip()
                        extracted.append(clean)
                
                elif isinstance(item, list):
                    # Если вложенный список, рекурсивно вызываем сами себя (плоская структура)
                    extracted.extend(cls.robust_query_extractor(item))

            # 3. Финальная фильтрация и дедупликация
            # Убираем пустые, слишком короткие строки и дубли
            unique_queries = []
            seen = set()
            for q in extracted:
                q_clean = q.strip('"\'').strip()
                if len(q_clean) > 3 and q_clean not in seen:
                    unique_queries.append(q_clean)
                    seen.add(q_clean)

            # 4. Если ничего не нашли, но вход был словарём, пробуем вернуть ключи (иногда модель путает ключи и значения)
            if not unique_queries and isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                 for item in v:
                     unique_queries.extend([str(k) for k in item.keys() if len(str(k)) > 10])

            return unique_queries[:7] # Ограничиваем топ-7, чтобы не спамить
        
    def conduct_chapter_focused_research(self, master_plan: ArticleMasterPlan):
        """
        Проводит исследование, итеративно проходя по каждой главе плана, 
        чтобы найти факты, подтверждающие тезисы главы.
        """
        self.log_thought("INIT", f"Запуск целевого исследования для '{master_plan.main_title}'")
        
        processed_urls = set()
        
        # Проходим по каждой главе из детального плана
        for i, chapter in enumerate(master_plan.chapters):
            self.log_thought("CHAPTER_FOCUS", f"Поиск доказательств для Главы {i+1}: '{chapter.title}'")

            # --- 1. Генерация поисковых запросов на основе ТЗ главы ---

            num_queries = self.llm.config.QUERIES_PER_CHAPTER
            
            # Подготовка переменных для промпта
            topic = master_plan.main_title
            subtopic = chapter.title

            query_gen_prompt = (
                f"ГЛАВНАЯ ТЕМА: '{topic}'\n"
                f"ТЕМА ГЛАВЫ: '{subtopic}'\n"
                f"АСПЕКТЫ: {', '.join(chapter.key_points)}\n\n"
                
                f"ЗАДАЧА: Сформулируй {num_queries} поисковых запроса на РУССКОМ ЯЗЫКЕ.\n\n"
                
                f"ПРАВИЛА ЭФФЕКТИВНОГО ПОИСКА:\n"
                f"1.  **КРАТКОСТЬ:** Поисковики плохо понимают длинные предложения. Длина запроса — строго 3-6 слов. Никаких лишних предлогов.\n"
                f"2.  **ПРИВЯЗКА:** В каждом запросе ОБЯЗАТЕЛЬНО должно быть главное слово из темы: '{topic}'. Без него мы найдем информацию не из той индустрии.\n"
                f"3.  **ФОРМАТ ЗАПРОСОВ (SEO-Style):**\n"
                f"    - Статистика: '{topic} {subtopic} статистика 2025'\n"
                f"    - Кейсы: '{topic} примеры {subtopic}'\n"
                f"    - Аналитика: '{topic} {subtopic} проблемы решения'\n\n"

                f"ВЕРНИ ТОЛЬКО JSON-СПИСОК СТРОК."
            )
            
            query_plan = self.llm.generate_structured(
                system_prompt=PromptsLibrary.QUERY_GENERATOR_PROTOCOL,
                user_prompt=query_gen_prompt,
                response_model=self.SearchQueryPlan,
                complexity=ComplexityLevel.ANALYTICAL,
                temperature=0.5
            )
            
            # Используем название главы как резервный запрос, если генерация упала
            queries_to_run = [] 
            if query_plan and hasattr(query_plan, 'queries') and query_plan.queries:
                queries_to_run.extend(query_plan.queries)
            
            if not queries_to_run:
                queries_to_run = [chapter.title] 
                self.log_thought("WARN", "Не удалось сгенерировать запросы. Использую название главы как фоллбэк.")

            # --- 2. Выполнение запросов и анализ источников ---
            for query in queries_to_run:
                self.log_thought("HUNT", f"Выполняю запрос: '{query}'")
                search_results = self.tools.web.search_google_simulated(query, limit=self.llm.config.SEARCH_NUMBER)
                
                for item in search_results:
                    if item.url in processed_urls: continue
                    
                    content = self.tools.web.get_page_content(item.url)
                    if not content or len(content) < 500: continue
                    
                    processed_urls.add(item.url)
                    
                    # Передаем не только контент, но и контекст главы для точного анализа
                    facts = self._extract_facts_with_context(content, chapter)
                    
                    if facts:
                        # Сохраняем факты в граф с привязкой к главе
                        self.tools.graph.add_node(
                            f"ИСТОЧНИК (для главы '{chapter.title}'): {item.title}\nURL: {item.url}\n\n{facts}",
                            "verified_evidence", {"url": item.url, "title": item.title, "chapter_focus": chapter.title}
                        )
                        self.log_thought("SUCCESS", f"Найдены релевантные факты в '{item.title[:40]}...'")
        
        self.log_thought("DONE", f"Исследование завершено. Проанализировано {len(processed_urls)} уникальных источников.")
        self.tools.graph.save()

    def _extract_facts_with_context(self, content: str, chapter: ChapterBlueprint) -> str:
        """
        Извлекает факты, используя промпт EVIDENCE_DISTILLER_PROTOCOL и контекст главы.
        Улучшенная версия: более устойчива к ответам модели.
        """
        # Обрезаем контент, чтобы не превысить лимит контекста
        content_for_analysis = content[:12000]
        
        extraction_prompt = (
            f"КОНТЕКСТ ЗАДАЧИ:\n"
            f"- Тезис главы, который мы доказываем: '{chapter.core_thesis}'\n\n"
            f"ТЕКСТ ИСТОЧНИКА:\n---\n{content_for_analysis}\n---\n\n"
            f"ПРОТОКОЛ 'ПОЛНОЕ ПОГРУЖЕНИЕ':\n"
            f"1.  **ЦЕЛЬ:** Наша цель — не просто найти пару цифр, а полностью понять и извлечь логическую цепочку и доказательную базу автора источника. Представь, что тебе нужно составить полное досье, на основе которого другой человек напишет экспертный пересказ.\n\n"
            f"2.  **ДЕЙСТВИЕ:** Внимательно прочитай текст и извлеки из него ВСЮ информацию, относящуюся к нашей теме. Сгруппируй найденное по следующим категориям:\n\n"
            f"   - **ОСНОВНЫЕ УТВЕРЖДЕНИЯ:** (Главные мысли и выводы автора текста)\n"
            f"   - **ДОКАЗАТЕЛЬСТВА:** (Все цифры, статистика, проценты, результаты исследований)\n"
            f"   - **ПРИМЕРЫ:** (Названия компаний, проектов, технологий, конкретные кейсы)\n"
            f"   - **КОНЦЕПЦИИ:** (Определения важных терминов)\n\n"
            f"3.  **ПРАВИЛО:** Будь максимально подробным. Копируй целые предложения, если они важны. Лучше дать больше информации, чем упустить ключевую деталь. НЕ ДЕЛАЙ СОБСТВЕННЫХ ВЫВОДОВ, только извлекай.\n\n"
            f"4.  **ИСКЛЮЧЕНИЕ:** Если текст вообще не по теме, верни только одну строку: [RELEVANCE: LOW]."
        )
        # Используем основной системный промпт агента (EVIDENCE_DISTILLER_PROTOCOL)
        facts = self.llm.generate_text(
            self.role_prompt, 
            extraction_prompt, 
            ComplexityLevel.ANALYTICAL, 
            temperature=0.8 
        )

        # === УЛУЧШЕННАЯ ПРОВЕРКА ===
        # Мы отбрасываем результат, ТОЛЬКО если он ЯВНО помечен как нерелевантный
        # или если он слишком короткий, чтобы быть полезным.
        if facts and "RELEVANCE: LOW" not in facts.upper() and len(facts) > 50:
            self.log_thought("SUCCESS", f"Извлечены факты из источника ({len(facts)} симв.).")
            return facts
        else:
            # Логируем причину отказа
            if "RELEVANCE: LOW" in facts.upper():
                self.log_thought("REJECTED", f"Источник отброшен как нерелевантный.")
            else:
                self.log_thought("REJECTED", f"Извлеченный текст слишком короткий или пустой.")
            return ""


    def conduct_broad_research(self, master_plan: ArticleMasterPlan):
        """
        Фаза 1: Широкое сканирование. Собирает общий пул знаний по всей теме.
        """
        self.log_thought("BROAD_SCAN_INIT", f"Запуск широкого сканирования по теме: '{master_plan.main_title}'")

        # 1. Генерируем 3-4 общих, но мощных запроса на основе всего плана
        query_gen_prompt = (
            f"ПЛАН ДОКУМЕНТА:\n"
            f"- Заголовок: {master_plan.main_title}\n"
            f"- Основная цель: {master_plan.abstract_objective}\n\n"
            f"ЗАДАЧА: Сформулируй 3-4 наиболее важных и всеобъемлющих поисковых запроса для сбора фундаментальной информации по этой теме. "
            f"Запросы должны охватывать ключевые концепции, статистику и основные примеры."
        )
        query_plan = self.llm.generate_structured(
            system_prompt=PromptsLibrary.QUERY_GENERATOR_PROTOCOL,
            user_prompt=query_gen_prompt,
            response_model=self.SearchQueryPlan,
            complexity=ComplexityLevel.ANALYTICAL
        )

        if not query_plan or not query_plan.queries:
            self.log_thought("WARN", "Не удалось сгенерировать общие запросы. Использую заголовок.")
            queries_to_run = [master_plan.main_title]
        else:
            queries_to_run = query_plan.queries

        # 2. Выполняем поиск и сохраняем ВСЕ найденное в граф
        self._execute_search_and_save(queries_to_run, "broad_scan_source")
        self.log_thought("BROAD_SCAN_DONE", "Широкое сканирование завершено. База знаний наполнена.")
        self.tools.graph.save()


    def _execute_search_and_save(self, queries: List[str], node_type_prefix: str):
        """
        Вспомогательный метод для выполнения поиска и сохранения результатов в граф.
        Версия 2.0: Отказоустойчивая, без дублирования кода.
        """
        processed_urls = set()
        for query in queries:
            self.log_thought("HUNT", f"Выполняю запрос: '{query}'")
            # Предполагается, что search_google_simulated не вызовет критическую ошибку,
            # а вернет пустой список в случае сбоя.
            search_results = self.tools.web.search_google_simulated(query, limit=self.llm.config.SEARCH_NUMBER)

            for item in search_results:
                if item.url in processed_urls:
                    continue
                
                # Обрабатываем каждый источник в отдельном блоке try-except.
                # Это гарантирует, что ошибка в одном источнике не остановит весь цикл.
                try:
                    processed_urls.add(item.url) # Добавляем URL здесь, чтобы не пытаться повторно при ошибке

                    content = self.tools.web.get_page_content(item.url)
                    if not content or len(content) < 300:
                        # Логируем пропуск, чтобы было понятно, почему источник не в графе
                        self.log_thought("SKIP", f"Источник '{item.title[:40]}...' пропущен (недостаточно контента).")
                        continue

                    # Используем базовый экстрактор, так как у нас еще нет контекста главы
                    extraction_prompt = (
                        f"ТЕКСТ ИСТОЧНИКА:\n---\n{content[:12000]}\n---\n\n"
                        f"ЗАДАЧА: Извлеки из этого текста ключевые факты, цифры и основные тезисы. Составь краткое, но информативное досье."
                    )
                    facts = self.llm.generate_text(self.role_prompt, extraction_prompt, ComplexityLevel.ANALYTICAL)

                    if facts and "RELEVANCE: LOW" not in facts.upper():
                        self.tools.graph.add_node(
                            content=f"ИСТОЧНИК: {item.title}\nURL: {item.url}\n\n{facts}",
                            node_type=f"{node_type_prefix}",
                            meta={"url": item.url, "title": item.title, "query_source": query}
                        )
                        self.log_thought("SUCCESS", f"Найден и обработан источник: '{item.title[:40]}...'")
                    else:
                        self.log_thought("REJECTED", f"Источник '{item.title[:40]}...' отклонен LLM как нерелевантный.")

                except Exception as e:
                    # Логируем ошибку с указанием URL для легкой отладки
                    self.log_thought("ERROR", f"Не удалось обработать источник {item.url}: {e}")
                    # `continue` критически важен: он переходит к следующему item, не прерывая цикл
                    continue
# ==============================================================================
# 18. ПИСАТЕЛЬ (THE WRITER) - ГЕНЕРАЦИЯ КОНТЕНТА + ВИЗУАЛИЗАЦИЯ 
# ==============================================================================

class WriterAgent(BaseAgent):
    """
    Отвечает за написание текста. Версия v8.3 (Configured).
    """
    
    def __init__(self, engine, tools):
        super().__init__("Writer", PromptsLibrary.WRITER_PHASE_1_ANALYSIS, engine, tools)
        
        cfg = self.llm.config  # Короткая ссылка на конфиг
        
        self.EVIDENCE_CHUNK_SIZE = cfg.WRITER_EVIDENCE_CHUNK_SIZE
        self.MAX_CONTEXT_CHARS = cfg.WRITER_MAX_CONTEXT_CHARS
        self.TEMP_CREATIVE = cfg.TEMP_CREATIVE
        self.MIN_TEXT_LENGTH = cfg.WRITER_MIN_TEXT_LENGTH

    def draft_section(self, section: ChapterBlueprint, prev_context_summary: str, next_chapter_title: str, evidence_block: str) -> str:
        """
        Создает черновик главы, используя RAG для доступа к "долгосрочной памяти" (графу знаний).
        """
        try:
            self.log_thought("INIT", f"Начинаю работу над главой: '{section.title}'")

            # === 1. АКТИВНЫЙ ПОИСК ВНУТРЕННЕГО КОНТЕКСТА (RAG) ===
            self.log_thought("SELF_REFLECTION", "Формирую запросы к долгосрочной памяти...")
            context_query_prompt = (
                f"Я собираюсь написать главу на тему '{section.title}', ключевой тезис которой: '{section.core_thesis}'.\n"
                f"Чтобы обеспечить логическую связь и избежать повторений, какие ключевые факты, выводы или цифры из предыдущих глав мне необходимо вспомнить и учесть?\n"
                f"Сформулируй 1-2 кратких поисковых запроса для внутренней базы знаний (графа)."
            )
            context_queries_raw = self.llm.generate_text(
                "Ты - ассистент, помогающий автору работать с его базой знаний.",
                context_query_prompt,
                ComplexityLevel.ROUTINE
            )
            context_queries = [q.strip("-* ").strip() for q in context_queries_raw.split('\n') if q.strip()]

            retrieved_context = ""
            if context_queries:
                self.log_thought("MEMORY_SEARCH", f"Запросы к памяти: {context_queries}")
                for q in context_queries:
                    search_results = self.tools.graph.search_text(q)
                    if search_results:
                        retrieved_context += f"СПРАВКА ИЗ ПАМЯТИ ПО ЗАПРОСУ '{q}':\n{search_results[0]['content'][:1500]}...\n\n"

            if not retrieved_context:
                retrieved_context = "Внутренний поиск не дал релевантных результатов. Опирайся на общую сводку."

            # --- 3. ПОДГОТОВКА ФИНАЛЬНОГО ПРОМПТА ---
            # Проверка типа key_points (защита от сбоев)
            if isinstance(section.key_points, list):
                formatted_key_points = "- " + "\n- ".join(section.key_points)
            else:
                formatted_key_points = str(section.key_points)

            writing_prompt = (
                f"=== 1. ОБЩИЙ КОНТЕКСТ (КРАТКАЯ ПАМЯТЬ) ===\n"
                f"Сводка предыдущих глав:\n...{prev_context_summary[-self.MAX_CONTEXT_CHARS:]}\n\n"

                f"=== 2. СПРАВКА ИЗ ДОЛГОСРОЧНОЙ ПАМЯТИ (RAG) ===\n"
                f"{retrieved_context}\n\n"

                f"=== 3. ВНЕШНИЕ ФАКТЫ ДЛЯ ЭТОЙ ГЛАВЫ (ИССЛЕДОВАНИЕ) ===\n"
                f"{evidence_block}\n\n"

                f"=== 4. ТВОЕ ТЕХНИЧЕСКОЕ ЗАДАНИЕ НА ГЛАВУ ===\n"
                f"ЗАГОЛОВОК: {section.title}\n"
                f"ГЛАВНЫЙ ТЕЗИС (что нужно доказать): {section.core_thesis}\n"
                f"КЛЮЧЕВЫЕ ПУНКТЫ ДЛЯ РАСКРЫТИЯ:\n{formatted_key_points}\n"
                f"ЛОГИЧЕСКИЙ МОСТ (к чему подвести в конце): {section.narrative_link}\n\n"

                f"=== ЗАДАЧА ===\n"
                f"Напиши ПОЛНЫЙ, ДЕТАЛЬНЫЙ и РАЗВЕРНУТЫЙ текст главы. Используй ВСЕ предоставленные контексты.\n\n"
                f"ТРЕБОВАНИЯ К ТЕКСТУ:\n"
                f"- **ГЛУБИНА:** Не просто перечисляй факты, а объясняй их, строй логические цепочки, делай выводы.\n"
                f"- **ОБЪЕМ:** Текст должен быть исчерпывающим. Минимального объёма нет. Пиши столько, сколько нужно для полного раскрытия темы, не экономь слова.\n\n"
                f"*** ВАЖНО: ЗАПРЕТ НА СКВОЗНУЮ НУМЕРАЦИЮ ***\n"
                f"Если ты создаешь нумерованные списки, ВСЕГДА начинай их с '1.'.\n"
                f"Строго следуй ТЗ."
            )

            final_text = self.llm.generate_text(
                system_prompt=PromptsLibrary.WRITER_PHASE_2_EXECUTION,
                user_prompt=writing_prompt,
                complexity=ComplexityLevel.CREATIVE,
                temperature=self.TEMP_CREATIVE
            )

            # --- 4. ОЧИСТКА И ПРОВЕРКА ---
            clean_text = self.llm._clean_llm_output(final_text).strip()
            # Убираем заголовок главы, если модель его продублировала
            title_pattern = re.compile(fr"^[\#\*]*\s*{re.escape(section.title)}\s*[\#\*]*[:\.]?\s*", re.IGNORECASE)
            clean_text = title_pattern.sub("", clean_text)

            if not clean_text or len(clean_text) < self.MIN_TEXT_LENGTH:
                self.log_thought("ERROR", "Сбой генерации (текст слишком короткий).")
                return f"*[Ошибка генерации главы '{section.title}'. Требуется рерайт.]*"

            return clean_text

        except Exception as e:
            self.logger.error(f"WriterAgent Critical Error in draft_section: {e}", exc_info=True)
            return f"*[Критическая системная ошибка при генерации главы '{section.title}']*"

    def _gather_evidence_for_chapter(self, section: ChapterBlueprint) -> Dict[str, Any]:
        """
        Собирает факты из Графа Знаний.
        === ИСПРАВЛЕНИЕ 2: Возвращает Словарь, а не Строку ===
        """
        empty_result = {"text": "Факты не найдены.", "source_ids": []}

        if not hasattr(self.tools, 'graph') or not self.tools.graph:
            return empty_result

        # 1. Поиск по ключевым словам
        search_query = f"{section.title} {section.core_thesis}"
        
        try:
            all_nodes = self.tools.graph.search_text(search_query)
        except Exception as e:
            self.logger.warning(f"Graph search failed: {e}")
            all_nodes = []
        
        evidence_nodes_content = []
        evidence_nodes_ids = []
        
        for n in all_nodes:
            content = n.get('content')
            node_id = n.get('id')
            if content and node_id:
                # Фильтруем, берем только источники или проверенные факты
                if 'ИСТОЧНИК:' in content or n.get('type') in ['verified_evidence', 'broad_scan_source', 'document']:
                    evidence_nodes_content.append(content)
                    evidence_nodes_ids.append(node_id)
        
        full_evidence_text = "\n\n---\n\n".join(evidence_nodes_content)
        
        # Если ничего не нашли
        if not full_evidence_text:
            return {"text": "Специфические факты отсутствуют. Опирайся на общую логику.", "source_ids": []}

        # 2. УМНАЯ ФИЛЬТРАЦИЯ (Smart Chunking)
        final_text = full_evidence_text
        
        if len(full_evidence_text) > self.EVIDENCE_CHUNK_SIZE:
            self.log_thought("FILTERING", f"Сжатие данных ({len(full_evidence_text)} симв)...")
            
            chunks = [full_evidence_text[i:i + self.EVIDENCE_CHUNK_SIZE] for i in range(0, len(full_evidence_text), self.EVIDENCE_CHUNK_SIZE)]
            filtered_chunks = []
            
            filtering_prompt_template = (
                f"ЦЕЛЬ ГЛАВЫ: '{section.title}'. Основной тезис: {section.core_thesis}\n"
                f"ФРАГМЕНТ ДОСЬЕ:\n{{text}}\n\n"
                f"ЗАДАЧА: Выбери из этого фрагмента ТОЛЬКО факты (цифры, имена, кейсы), которые критически важны для цели главы.\n"
                f"Игнорируй воду. Сохрани исходные формулировки."
            )
        
            for i, chunk in enumerate(chunks):
                if i > 3: break # Лимит чанков чтобы не зависнуть
                summary = self.llm.generate_text(
                    system_prompt="Ты — аналитик данных.",
                    user_prompt=filtering_prompt_template.format(text=chunk),
                    complexity=ComplexityLevel.ANALYTICAL,
                    temperature=0.8
                )
                filtered_chunks.append(summary)
            
            final_text = "\n\n".join(filtered_chunks)

        return {"text": final_text, "source_ids": evidence_nodes_ids}

    def revise_text(self, original_text: str, feedback: ReviewFeedback) -> str:
        """
        Исправляет текст на основе замечаний Критика.
        """
        # Если правок нет, возвращаем оригинал
        if not feedback.required_edits: 
            return original_text

        self.log_thought("REVISING", "Вношу правки согласно замечаниям критика...")

        formatted_edits = "- " + "\n- ".join(feedback.required_edits)
        
        revision_prompt = (
            f"--- SYSTEM DIRECTIVE: EDITORIAL REVISION PROTOCOL ---\n\n"
            f"РОЛЬ: Вы — Главный Редактор и Эксперт предметной области. Ваша цель — довести черновик до совершенства.\n\n"
            
            f"=== ВХОДНЫЕ ДАННЫЕ: ЧЕРНОВИК ===\n"
            f"{original_text}\n"
            f"================================\n\n"
            
            f"=== ОТЧЕТ КРИТИКА (ТРЕБУЕМЫЕ ПРАВКИ) ===\n"
            f"{formatted_edits}\n"
            f"========================================\n\n"
            
            f"ЗАДАЧА: Проведи хирургически точную редактуру текста, интегрируя все указанные правки.\n\n"
            
            f"СТРОГИЙ ОПЕРАЦИОННЫЙ ПРОТОКОЛ:\n"
            f"1.  **ПОЛНОТА ИНТЕГРАЦИИ:** Каждое замечание из списка должно быть отработано. Игнорирование правок запрещено.\n"
            f"2.  **ЗАПРЕТ НА СОКРАЩЕНИЕ:** Строжайше запрещено уменьшать объем текста или удалять детали, которые не требовали исправления. Текст должен остаться таким же глубоким и подробным. Разрешено только дополнять.\n"
            f"3.  **СОХРАНЕНИЕ СТРУКТУРЫ:** Ты обязан сохранить всю Markdown-разметку (заголовки #, жирный шрифт **, списки -, цитаты >). Визуальная структура документа должна остаться идеальной.\n"
            f"4.  **СТИЛИСТИЧЕСКАЯ ЦЕЛОСТНОСТЬ:** Сохрани авторский, аналитический и экспертный тон повествования. Текст не должен звучать как машинный перевод.\n"
            f"5.  **БЕСШОВНОСТЬ:** Новые вставки должны органично вплетаться в повествование. Не должно быть резких скачков логики.\n\n"
            
            f"ФОРМАТ ВЫВОДА:\n"
            f"Верни ТОЛЬКО полный отредактированный текст главы. Никаких вступлений вроде 'Вот исправленный вариант', никаких комментариев. Сразу начинай с заголовка или текста."
        )

        # Используем чуть меньшую температуру (0.5), чтобы модель была точнее в исправлениях
        return self.llm.generate_text(
            system_prompt=self.role_prompt, 
            user_prompt=revision_prompt, 
            complexity=ComplexityLevel.CREATIVE, 
            temperature=0.5
        )


    def generate_visual_content(self, section_text: str, section_title: str) -> Optional[str]:
        """
        Генератор графиков v9.2 (Style-Matched & Fail-Safe).
        
        Особенности:
        1. Полная стилизация под DocumentForge (Times New Roman, размеры шрифтов).
        2. "Золотое правило" надежности: запрет на парсинг текста кодом.
        3. Защита от обрезания графиков (tight_layout).
        """
        # 1. Проверка на наличие числовых данных
        if not re.search(r'\d', section_text):
            self.log_thought("VISUAL_SKIP", f"В разделе '{section_title}' нет числовых данных для визуализации.")
            return None

        self.log_thought("VISUAL_INIT", f"Запуск цикла генерации/отладки для '{section_title}'...")

        # 2. Подготовка путей и уникальных имен
        scripts_dir = os.path.join(CONFIG.WORK_DIR, "temp_scripts")
        output_dir = os.path.join(CONFIG.WORK_DIR, "visuals")
        os.makedirs(scripts_dir, exist_ok=True)
        os.makedirs(output_dir, exist_ok=True)

        unique_id = f"{int(time.time())}_{uuid.uuid4().hex[:6]}"
        chart_filename = f"chart_{unique_id}.png"
        script_filename = f"script_{unique_id}.py"
        script_filepath = os.path.join(scripts_dir, script_filename)
        chart_filepath = os.path.join(output_dir, chart_filename)

        # 3. Настройка стиля под DocumentForge (Times New Roman)
        # Мы используем базовый стиль (например, ggplot) для сетки, 
        # но поверх него накладываем шрифты документа.
        base_styles = ['ggplot', 'seaborn-v0_8-whitegrid', 'bmh']
        chosen_style = random.choice(base_styles)
        chosen_palette = random.choice(['viridis', 'plasma', 'coolwarm', 'magma', 'rocket', 'mako'])

        # Параметры, копирующие стиль вашего DOCX
        docx_style_params = {
            'font.family': 'serif',
            'font.serif': ['Times New Roman'],
            'font.size': 11,              # Основной текст как в Normal
            'axes.titlesize': 14,         # Заголовок чуть меньше Heading 1
            'axes.titleweight': 'bold',
            'axes.labelsize': 12,         # Подписи осей
            'xtick.labelsize': 10,
            'ytick.labelsize': 10,
            'text.color': 'black',
            'axes.labelcolor': 'black',
            'xtick.color': 'black',
            'ytick.color': 'black',
            'legend.fontsize': 10,
            'figure.titlesize': 14
        }
        params_str = str(docx_style_params)

        # 4. Формирование усиленного промпта
        base_prompt = (
            f"--- ЗАДАЧА ГЕНЕРАЦИИ КОДА: ВИЗУАЛИЗАЦИЯ ДАННЫХ ---\n\n"
            f"**РОЛЬ:** Вы — Senior Python-разработчик. Ваша цель — создать надежный код для аналитического отчета в строгом корпоративном стиле.\n\n"
            
            f"**КОНТЕКСТ:**\n"
            f"- **Тема:** '{section_title}'.\n"
            f"- **Данные:**\n"
            f"  ---\n{section_text[:5000]}\n---\n\n"
            
            f"**ГЛАВНАЯ ЦЕЛЬ:** Сгенерировать Python-скрипт для создания **простого** графика. Стиль должен идеально сочетаться с документом Word (Times New Roman).\n\n"
            
            f"--- АЛГОРИТМ РАБОТЫ (СТРОГО СОБЛЮДАТЬ) ---\n\n"
            f"**ШАГ 1: Подготовка данных (ЗОЛОТОЕ ПРАВИЛО):**\n"
            f"   - **ЗАПРЕЩЕНО** писать код, который ищет цифры в тексте (regex, split и т.д.). Это вызывает ошибки.\n"
            f"   - **ОБЯЗАТЕЛЬНО** сам прочитай текст, выбери данные и **создай DataFrame вручную** в коде.\n"
            f"   - *Правильно:* `df = pd.DataFrame({{'год': [2020, 2021], 'выручка': [100, 150]}})`\n"
            f"   - *Неправильно:* `df = pd.read_csv(...)` или `data = re.findall(...)`\n\n"

            f"**ШАГ 2: Визуализация:**\n"
            f"   - Используй `import matplotlib.pyplot as plt` и `import seaborn as sns`.\n"
            f"   - Создай фигуру: `plt.figure(figsize=(10, 6))`.\n"
            f"   - Примени стиль: `plt.style.use('{chosen_style}')`.\n"
            f"   - **ВАЖНО:** Примени настройки шрифта из документа: `plt.rcParams.update({params_str})`.\n"
            f"   - Построй график (barplot, lineplot или pie chart). Не используй сложные комбинации.\n\n"

            f"**ШАГ 3: Оформление:**\n"
            f"   - Добавь заголовок `plt.title(...)` и подписи осей.\n"
            f"   - **ЗАПРЕЩЕНО** использовать двойные оси (`twinx`).\n"
            f"   - Если нужно добавить примечание (например, 'Рост +5%'), используй `plt.figtext(0.99, 0.01, '...', ha='right')`, а не аннотации на графике.\n\n"

            f"**ШАГ 4: Сохранение:**\n"
            f"   - `plt.tight_layout()` (Обязательно, чтобы не обрезать текст).\n"
            f"   - `output_path = r'{chart_filepath}'` (Использовать эту переменную).\n"
            f"   - `plt.savefig(output_path, dpi=300, bbox_inches='tight')`.\n\n"

            f"**СПЕЦИФИКАЦИЯ ВЫВОДА:**\n"
            f"Верни ТОЛЬКО валидный код Python. Начинай сразу с `import pandas as pd`."
        )

        # 5. Цикл самоотладки
        error_feedback = ""
        for attempt in range(1, 4):
            self.log_thought("VISUAL_ATTEMPT", f"Попытка {attempt}/3...")
            
            current_prompt = f"{error_feedback}{base_prompt}"
            
            code_response = self.llm.generate_text(
                system_prompt="You are an expert Python developer. Write robust, error-free code for data visualization.",
                user_prompt=current_prompt,
                complexity=ComplexityLevel.ANALYTICAL
            )
            # Очистка от маркдауна
            code = code_response.strip().replace("```python", "").replace("```", "").strip()

            if not code or "import" not in code:
                self.log_thought("VISUAL_REJECT", f"Попытка {attempt} провалена. Нет кода в ответе.")
                error_feedback = self._create_error_feedback("No code generated.", code)
                continue

            try:
                # Добавляем "тихий" режим для matplotlib
                final_code = f"import matplotlib\nmatplotlib.use('Agg')\nimport warnings\nwarnings.filterwarnings('ignore')\n{code}"
                
                with open(script_filepath, 'w', encoding='utf-8') as f:
                    f.write(final_code)

                # WARNING: This executes AI-generated code locally. 
                # Ensure running in a sandboxed environment if used in production.

                # Запуск во внешнем процессе
                result = subprocess.run(
                    [sys.executable, script_filepath],
                    capture_output=True, text=True, timeout=30, encoding='utf-8'
                )

                # Проверка результата
                if os.path.exists(chart_filepath) and os.path.getsize(chart_filepath) > 1024:
                    self.log_thought("VISUAL_SUCCESS", f"График '{chart_filename}' успешно создан.")
                    return chart_filepath
                else:
                    error_log = result.stderr if result.stderr else result.stdout
                    self.log_thought("VISUAL_FAIL", f"Попытка {attempt} провалена. Ошибка: {error_log[:300]}")
                    error_feedback = self._create_error_feedback(error_log, code)

            except subprocess.TimeoutExpired:
                self.log_thought("VISUAL_FAIL", f"Попытка {attempt} провалена (Timeout).")
                error_feedback = self._create_error_feedback("Script timed out.", code)
            except Exception as e:
                self.log_thought("VISUAL_ERROR", f"Критическая ошибка: {e}")
                error_feedback = self._create_error_feedback(str(e), code)
        
        self.log_thought("VISUAL_GIVE_UP", f"Не удалось создать график для '{section_title}'.")
        return None

    def _create_error_feedback(self, error_log: str, faulty_code: str) -> str:
        """
        Вспомогательный метод для создания блока отладочной информации для LLM на русском языке.
        """
        safe_error_log = error_log[:2000]
        
        return (
            f"**ОТЛАДОЧНАЯ ИНФОРМАЦИЯ (ПРЕДЫДУЩАЯ ПОПЫТКА НЕ УДАЛАСЬ):**\n"
            f"Сгенерированный тобой Python-скрипт вызвал ошибку. Ты ОБЯЗАН проанализировать ошибку и свой код, чтобы исправить её.\n\n"
            
            f"**ТРЕЙСБЕК ОШИБКИ:**\n"
            f"```\n{safe_error_log}\n```\n\n"
            
            f"**ТВОЙ ОШИБОЧНЫЙ КОД:**\n"
            f"```python\n{faulty_code}\n```\n\n"
            
            f"**НОВАЯ ЗАДАЧА:** Проанализируй ошибку. Скорее всего, ты нарушил одно из правил из раздела 'КРАСНЫЕ ЛИНИИ'. "
            f"**ПРОВЕРЬ, НЕ ИСПОЛЬЗОВАЛ ЛИ ТЫ ЧИСЛО В `hue` ИЛИ `twinx()`/`twiny()`?**\n"
            f"Верни новую, полностью исправленную и **упрощенную** версию скрипта. Не повторяй ту же ошибку. Сгенерируй полностью готовый к запуску скрипт с нуля.\n\n"
        )

class StylePolymath(BaseAgent):
    """
    Агент-"Душа". Занимается исключительно стилизацией и "очеловечиванием".
    """
    def __init__(self, engine, tools):
        super().__init__("StylePolymath", PromptsLibrary.STYLE_POLYMATH_DIRECTIVE, engine, tools)

    def humanize_text(self, draft_text: str, context_mood: str = "Analytical") -> str:
        self.log_thought("FEELING", "Вдыхаю жизнь в текст (Humanizing)...")
        
        # Если текст слишком большой, бьем на части, чтобы не потерять качество (для 4B это критично)
        if len(draft_text) > 8000:
            self.log_thought("WARN", "Текст слишком велик. Стилизую частями.")
            # Упрощенная логика: стилизуем начало и конец, середину оставляем (чтобы не сломать)
            # В идеале нужен сплиттер, но для надежности пока так:
            to_process = draft_text[:4000]
            rest = draft_text[4000:]
        else:
            to_process = draft_text
            rest = ""

        humanize_prompt = (
            f"ЧЕРНОВИК:\n{to_process}\n\n"
            f"ЗАДАЧА: Перепиши этот текст, следуя 'Протоколу Очеловечивания'.\n"
            f"Сделай его менее 'машинным' и более 'авторским'.\n"
            f"Сохрани все факты и цифры.\n"
            f"ВЕРНИ ТОЛЬКО ГОТОВЫЙ ТЕКСТ."
        )

        polished_text = self.llm.generate_text(
            system_prompt=self.role_prompt,
            user_prompt=humanize_prompt,
            complexity=ComplexityLevel.CREATIVE, # Включаем креативный режим на максимум
            temperature=0.8
        )

        # Защита: если модель вернула слишком короткий текст (сбой), возвращаем оригинал
        if len(polished_text) < len(to_process) * 0.5:
            self.log_thought("FAIL", "Стилизация не удалась (текст обрезан). Откат к черновику.")
            return draft_text

        final_result = polished_text + rest
        
        # Удаляем возможные артефакты ("Вот переписанный текст:")
        final_result = re.sub(r'^(Вот|Here is).*?:\n', '', final_result, flags=re.IGNORECASE).strip()
        
        return final_result

class ContinuityDirector(BaseAgent):
    """
    Агент-"Мост". Следит, чтобы главы не выглядели разрозненными кусками.
    """
    def __init__(self, engine, tools):
        super().__init__("Continuity", PromptsLibrary.CONTINUITY_DIRECTIVE, engine, tools)

    def bridge_chapters(self, prev_chapter_end: str, current_chapter_text: str) -> str:
        if not prev_chapter_end:
            return current_chapter_text # Первая глава, связывать не с чем

        self.log_thought("BRIDGING", "Строю нарративный мост между главами...")
        
        # Выделяем первый абзац текущей главы
        parts = current_chapter_text.split('\n\n')
        first_paragraph = parts[0]
        rest_of_chapter = '\n\n'.join(parts[1:])
        
        bridge_prompt = (
            f"КОНЦОВКА ПРЕДЫДУЩЕЙ ГЛАВЫ:\n...{prev_chapter_end[-2000:]}\n\n"
            f"НАЧАЛО ТЕКУЩЕЙ ГЛАВЫ:\n{first_paragraph}\n\n"
            f"ЗАДАЧА: Перепиши НАЧАЛО текущей главы так, чтобы оно логически вытекало из концовки предыдущей.\n"
            f"Используй связку (например: 'Этот парадокс подводит нас к...', 'Однако, как мы выяснили...').\n"
            f"Верни ТОЛЬКО один новый абзац."
        )

        new_intro = self.llm.generate_text(
            system_prompt=self.role_prompt,
            user_prompt=bridge_prompt,
            complexity=ComplexityLevel.CREATIVE,
            temperature=0.8
        )
        
        # Если генерация удалась, подменяем
        if len(new_intro) > 50:
             return f"{new_intro}\n\n{rest_of_chapter}"
        
        return current_chapter_text

# ==============================================================================
# 19. КРИТИК (THE CRITIC) 
# ==============================================================================

class CriticAgent(BaseAgent):
    """
    Проводит аудит текста на соответствие протоколам.
    """
    def __init__(self, engine, tools):
        super().__init__("Critic", PromptsLibrary.CRITIC_DIRECTIVE, engine, tools)

    def evaluate_draft(self, text: str, section_goal: str) -> ReviewFeedback:
        self.log_thought("AUDIT", "Провожу аудит методологии и фактов...")
        
        # === ИСПРАВЛЕНО: Используем переменную из конфига ===
        # Обрезаем текст до лимита символов, заданного в SystemConfig
        text_for_audit = text[:self.llm.config.CRITIC_INPUT_MAX_CHARS]
        
        eval_prompt = (
            f"ЦЕЛЬ РАЗДЕЛА: {section_goal}\n\n"
            f"ТЕКСТ НА АУДИТ:\n{text_for_audit}\n\n"
            f"ЗАДАЧА: Проверь текст по протоколу (Factual Verification, Logical Fallacies, Argument Strength).\n"
            f"Верни JSON-отчет с вердиктом и списком КОНКРЕТНЫХ правок."
        )
        
        feedback = self.llm.generate_structured(
            system_prompt=self.role_prompt,
            user_prompt=eval_prompt,
            response_model=ReviewFeedback,
            complexity=ComplexityLevel.ANALYTICAL
        )
        
        status = "ПРИНЯТО" if feedback.is_approved else "ОТКЛОНЕНО"
        self.log_thought("VERDICT", f"Вердикт: {status} (Оценка: {feedback.score}).")
        return feedback


class GhostwriterAgent(BaseAgent):
    """
    Агент-"Призрак". Финальный слой обработки, делающий текст неотличимым от человеческого.
    """
    def __init__(self, engine, tools):
        super().__init__("Ghostwriter", PromptsLibrary.GHOST_PROTOCOL, engine, tools)

    def make_undetectable(self, text: str) -> str:
        self.log_thought("GHOSTING", "Запуск протокола 'Призрак' (Anti-AI)...")

        if len(text) > 8000:
            chunks = [text[i:i+8000] for i in range(0, len(text), 8000)]
        else:
            chunks = [text]
        
        processed_chunks = []
        for chunk in chunks:
            ghost_prompt = (
                f"ТЕКСТ ДЛЯ ОБРАБОТКИ:\n{chunk}\n\n"
                f"ЗАДАЧА: Примени протокол 'Humanization & Obfuscation'.\n"
                f"Сделай 5-6 точечных изменений, чтобы текст выглядел написанным человеком."
            )
            
            processed_chunk = self.llm.generate_text(
                system_prompt=self.role_prompt,
                user_prompt=ghost_prompt,
                complexity=ComplexityLevel.CREATIVE,
                temperature=0.7 # Температура нужна для креативных синонимов
            )

            # Защита: если модель вернула пустой или сломанный текст, используем оригинал
            if len(processed_chunk) < len(chunk) * 0.8:
                self.log_thought("WARN", "Протокол 'Призрак' вернул слишком короткий текст. Откат к оригиналу.")
                processed_chunks.append(chunk)
            else:
                processed_chunks.append(processed_chunk)

        final_text = "\n\n".join(processed_chunks)
        self.log_thought("DONE", "Текст успешно 'очеловечен'.")
        return final_text


class TelegramBridge:
    """
    Информационный мост v2.0. Отправляет структурированные отчеты о ходе миссии.
    """
    def __init__(self, token: str, chat_id: str):
        self.token = token
        self.chat_id = chat_id
        self.enabled = bool(token and chat_id and "ТВОЙ" not in token and "ВСТАВЬ" not in chat_id)
        self.queue = queue.Queue()
        self.session = requests.Session()
        
        if self.enabled:
            threading.Thread(target=self._worker, daemon=True).start()
            logger.info("Telegram Bridge v2.0 активирован.")

    # --- Публичные методы для вызова из системы ---

    def send_mission_start(self, topic: str):
        """Сообщает о начале новой миссии и присылает исходный промпт."""
        self._queue_task("start", {"topic": topic})

    def send_plan(self, plan: ArticleMasterPlan):
        """Отправляет сгенерированный Архитектором план."""
        # Используем родной метод Pydantic вместо asdict
        self._queue_task("plan", {"plan": plan.model_dump()})

    def send_research_update(self, status: Literal["start", "done"], found_sources: int = 0):
        """Сообщает о начале или завершении фазы исследования."""
        self._queue_task("research", {"status": status, "found": found_sources})

    def send_writing_start(self):
        """Сообщает о начале фазы написания глав."""
        self._queue_task("writing_start")

    def send_chapter_update(self, current: int, total: int, title: str):
        """Отправляет прогресс-бар по написанию глав."""
        self._queue_task("chapter_done", {"current": current, "total": total, "title": title})
        
    def send_final_summary(self, title: str, file_path: str, elapsed_time: float):
        """Отправляет финальный отчет с документом и статистикой."""
        self._queue_task("final_summary", {"title": title, "path": file_path, "time": elapsed_time})
        
    def send_error(self, error_msg: str, traceback_info: str):
        """Отправляет сообщение о критической ошибке."""
        self._queue_task("error", {"error": error_msg, "traceback": traceback_info})

    # --- Внутренняя кухня (обработка очереди) ---

    def _queue_task(self, task_type: str, data: Dict = None):
        if not self.enabled: return
        self.queue.put({"type": task_type, "data": data or {}})

    def _worker(self):
        """Фоновый воркер, который отправляет сообщения, не блокируя основной поток."""
        while True:
            item = self.queue.get()
            try:
                task_type = item.get("type")
                data = item.get("data")
                
                if task_type == "start":
                    text = f"🚀 **НОВАЯ МИССИЯ ЗАПУЩЕНА**\n\n**Задача:**\n```\n{data['topic'][:1000]}...\n```"
                    self._send_message(text)
                
                elif task_type == "plan":
                    plan = data['plan']
                    # Формируем список глав
                    chapters_list = ""
                    for c in plan.get('chapters', []):
                        # Добавляем эмодзи книги для красоты
                        chapters_list += f"  📖 {c['title']}\n"
                    
                    # Обрезаем список глав, если он гигантский (защита от лимита ТГ)
                    if len(chapters_list) > 3000:
                        chapters_list = chapters_list[:3000] + "\n... [список сокращен]"

                    text = (
                        f"🏛️ **АРХИТЕКТОР ЗАВЕРШИЛ ПЛАНИРОВАНИЕ**\n\n"
                        f"**Название:** {plan.get('main_title', 'Без названия')}\n"
                        f"**Подзаголовок:** {plan.get('subtitle', 'Нет')}\n\n"
                        f"**Структура глав:**\n{chapters_list}"
                    )
                    self._send_message(text)

                elif task_type == "research":
                    if data['status'] == 'start':
                        text = "🔎 **ИССЛЕДОВАТЕЛЬ НАЧАЛ ПОИСК**\nСобираю данные из открытых источников..."
                    else:
                        text = f"✅ **ИССЛЕДОВАНИЕ ЗАВЕРШЕНО**\n\nНайдено и проанализировано **{data['found']}** релевантных источников. Данные загружены в Граф Знаний."
                    self._send_message(text)

                elif task_type == "writing_start":
                    text = "✍️ **ПИСАТЕЛЬ ПРИСТУПИЛ К РАБОТЕ**\nНачинаю генерацию текста глав..."
                    self._send_message(text)
                
                elif task_type == "chapter_done":
                    progress = "🟩" * data['current'] + "⬜️" * (data['total'] - data['current'])
                    text = (
                        f"**Прогресс: Глава {data['current']}/{data['total']} написана**\n\n"
                        f"[{progress}]\n\n"
                        f"*{data['title']}*"
                    )
                    self._send_message(text)
                
                elif task_type == "final_summary":
                    elapsed_min = data['time'] / 60
                    text = (
                        f"🎉 **МИССИЯ ВЫПОЛНЕНА!**\n\n"
                        f"**Итоговый документ:** {data['title']}\n"
                        f"**Затраченное время:** {elapsed_min:.1f} мин."
                    )
                    self._send_document(data['path'], text)

                elif task_type == "error":
                    text = (
                        f"❌ **КРИТИЧЕСКАЯ ОШИБКА**\n\n"
                        f"**Сообщение:** `{data['error']}`\n\n"
                        f"**Traceback:**\n```\n{data['traceback'][-1000:]}\n```"
                    )
                    self._send_message(text)

            except Exception as e:
                logger.error(f"Ошибка в TelegramBridge worker: {e}")
            finally:
                self.queue.task_done()

    def _send_message(self, text: str, disable_notification: bool = False):
        url = f"https://api.telegram.org/bot{self.token}/sendMessage"
        payload = {
            "chat_id": self.chat_id,
            "text": text,
            "parse_mode": "Markdown",
            "disable_notification": disable_notification,
        }
        try:
            self.session.post(url, json=payload, timeout=10)
        except Exception as e:
            logger.warning(f"Не удалось отправить сообщение в Telegram: {e}")

    def _send_document(self, filepath: str, caption: str):
        url = f"https://api.telegram.org/bot{self.token}/sendDocument"
        try:
            # Небольшая пауза, чтобы файловая система успела "отпустить" файл после записи
            time.sleep(1.0) 
            with open(filepath, 'rb') as f:
                self.session.post(
                    url, 
                    data={"chat_id": self.chat_id, "caption": caption, "parse_mode": "Markdown"}, 
                    files={"document": f},
                    timeout=60
                )
        except Exception as e:
            logger.error(f"Не удалось отправить документ в Telegram: {e}")

# ==============================================================================
# 20. класс-помощник 
# ==============================================================================

class ContextManager:
    """Управляет накоплением и сжатием контекста для LLM."""
    def __init__(self, initial_context: str, max_chars: int = 12000):
        self.full_log = [initial_context]
        self.summary_log = [initial_context]
        self.max_chars = max_chars

    def add(self, chapter_title: str, summary: str):
        entry = f"[Глава: {chapter_title}]: {summary}"
        self.full_log.append(entry)
        self.summary_log.append(entry)
        self._compress()

    def get(self) -> str:
        return "\n\n".join(self.summary_log)

    def _compress(self):
        """Если контекст слишком большой, 'забываем' детали старых глав."""
        current_length = sum(len(s) for s in self.summary_log)
        if current_length > self.max_chars:
            # Превращаем старые записи в одно предложение
            compressed_summary = f"...(Ранее обсуждались темы: {', '.join([s.split(':')[0][1:] for s in self.summary_log[1:-5]])})..."
            # Оставляем введение, сжатую середину и 5 последних глав
            self.summary_log = [self.summary_log[0], compressed_summary] + self.summary_log[-5:]


# ==============================================================================
# 21. ОРКЕСТРАТОР (GENESIS SYSTEM)
# ==============================================================================

class GenesisSystem:
    def __init__(self):
        self.config = CONFIG
        self.engine = LLMEngine(self.config)
        self.tools = GenesisToolbox()
        
        self.architect = ArchitectAgent(self.engine, self.tools)
        self.researcher = ResearcherAgent(self.engine, self.tools)
        self.writer = WriterAgent(self.engine, self.tools)
        self.style_agent = StylePolymath(self.engine, self.tools)
        self.continuity = ContinuityDirector(self.engine, self.tools)
        self.critic = CriticAgent(self.engine, self.tools)

        self.mission_log = []
        self.tg = TelegramBridge(self.config.TG_BOT_TOKEN, self.config.TG_CHAT_ID)
        
        # State Management
        self.state_dir = os.path.join(CONFIG.WORK_DIR, "states")
        os.makedirs(self.state_dir, exist_ok=True)
        self.state = None

    def log_system(self, msg: str):
        print(f"[SYSTEM] {msg}")
        self.mission_log.append(msg)
        logger.info(f"[SYSTEM] {msg}")

    def _get_state_path(self, topic: str) -> str:
        filename = hashlib.md5(topic.encode()).hexdigest() + ".json"
        return os.path.join(self.state_dir, filename)

    def _save_state(self):
        if not self.state: return
        path = self._get_state_path(self.state.topic)
        with open(path, 'w', encoding='utf-8') as f:
            f.write(self.state.to_json())
        logger.info(f"Mission State saved: {path}")

    def _load_state(self, topic: str) -> bool:
        path = self._get_state_path(topic)
        if os.path.exists(path):
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    self.state = MissionState.from_json(f.read())
                self.log_system(f"Восстановлено состояние: {self.state.status}, Глава {self.state.current_chapter_index}")
                return True
            except Exception as e:
                logger.error(f"Ошибка загрузки состояния: {e}")
        return False

    def produce_article(self, topic: str) -> str:
        """
        Главный оркестратор системы. Выполняет полный цикл создания документа.
        
        Версия 2.0 (Отказоустойчивая). Этот метод построен как конечный автомат (state machine),
        где каждый этап является транзакцией. Состояние сохраняется до и после каждого
        критического шага, обеспечивая возможность возобновления миссии после любого сбоя.
        """
        start_time = time.time()
        try:
            # ==========================================================================
            # ЭТАП 0: ИНИЦИАЛИЗАЦИЯ МИССИИ
            # ==========================================================================

            self.log_system("Protocol ENGAGED")

            if self._load_state(topic):
                self.log_system(f"Восстановление прерванной миссии: '{self.state.topic}'")
                self.log_system(f"Текущий статус: {self.state.status}")
            else:
                self.log_system(f"Запуск новой миссии: '{topic}'")
                self.state = MissionState(topic=topic, status="init")
                self._save_state() # Сохраняем самое начальное состояние

            self.tg.send_mission_start(self.state.topic)

            # ==========================================================================
            # ЭТАП 1: АРХИТЕКТУРА И РЕЦЕНЗИЯ ПЛАНА (Quality Gate #1)
            # ==========================================================================
            if self.state.status == "init":
                self.log_system("\n--- ЭТАП 1: ПРОЕКТИРОВАНИЕ СТРУКТУРЫ ---")
                
                master_plan = self.architect.create_master_plan(self.state.topic)
                if not master_plan or not master_plan.chapters:
                    raise RuntimeError("Критическая ошибка: Архитектор не смог создать базовый план.")
                
                self.log_system("Запуск архитектурной рецензии (Quality Gate #1)...")
                plan_review_prompt = (
                    f"ПЛАН ДОКУМЕНТА:\n{master_plan.model_dump_json(indent=2)}\n\n"
                    f"ЗАДАЧА: Ты — главный редактор. Оцени этот план по шкале от 1 до 10.\n"
                    f"КРИТЕРИИ: Логика, Глубина, Полнота. Если оценка < 7, план непригоден. Укажи, что исправить.\n"
                    f"Верни JSON: {{\"score\": <int>, \"is_approved\": <bool>, \"required_edits\": [\"...\"]}}"
                )
                plan_feedback = self.critic.llm.generate_structured(PromptsLibrary.CRITIC_DIRECTIVE, plan_review_prompt, ReviewFeedback, ComplexityLevel.ANALYTICAL)
                
                if not plan_feedback.is_approved and plan_feedback.score < 4:
                    raise RuntimeError(f"ПЛАН ОТКЛОНЕН (Оценка слишком низкая: {plan_feedback.score}). Причина: {plan_feedback.required_edits}")
                elif not plan_feedback.is_approved:
                    self.log_system(f"⚠️ Критик сделал замечания (Оценка: {plan_feedback.score}), но мы продолжаем работу.")
                self.log_system(f"✅ План прошел рецензию (Оценка: {plan_feedback.score}). Структура утверждена.")
                
                self.state.master_plan = master_plan.model_dump()
                self.state.status = "plan_approved"
                self._save_state()
                self.tg.send_plan(master_plan)
                try: logger.info(f"[STRUCTURE] {json.dumps([c.title for c in master_plan.chapters], ensure_ascii=False)}")
                except: pass
            
            master_plan = ArticleMasterPlan.model_validate(self.state.master_plan)

            # ==========================================================================
            # ЭТАП 2: ДВУХФАЗНОЕ ИССЛЕДОВАНИЕ
            # ==========================================================================
            if self.state.status == "plan_approved":
                self.log_system("\n--- ЭТАП 2.1: ШИРОКОЕ СКАНИРОВАНИЕ ТЕМЫ ---")
                self.tg.send_research_update(status="start")
                self.researcher.conduct_broad_research(master_plan)
                self.state.status = "broad_research_done"
                self._save_state()

            if self.state.status == "broad_research_done":
                self.log_system("\n--- ЭТАП 2.2: ГЛУБОКОЕ ПОГРУЖЕНИЕ ПО ГЛАВАМ ---")
                self.researcher.conduct_chapter_focused_research(master_plan)
                evidence_count = sum(1 for _, d in self.tools.graph.graph.nodes(data=True) if 'source' in d.get('type', ''))
                self.tg.send_research_update(status="done", found_sources=evidence_count)
                self.state.status = "research_done"
                self._save_state()

            # ==========================================================================
            # ЭТАП 3: ЦИКЛ НАПИСАНИЯ ГЛАВ
            # ==========================================================================

            if self.state.status in ["research_done", "writing_in_progress"]:
                self.log_system("\n--- ЭТАП 3: НАПИСАНИЕ И РЕДАКТУРА ГЛАВ ---")
                if self.state.status != "writing_in_progress":
                    self.tg.send_writing_start()
                    self.state.status = "writing_in_progress"
                    
                context_manager = ContextManager(initial_context=self.state.chapter_summaries[-1] if self.state.chapter_summaries else f"Начало документа '{master_plan.main_title}'.")

                for i, chapter_dict in enumerate(master_plan.chapters):
                    if i < self.state.current_chapter_index: continue
                    
                    chapter = ChapterBlueprint.model_validate(chapter_dict)
                    self.log_system(f"\n--- Глава {i+1}/{len(master_plan.chapters)}: '{chapter.title}' ---")
                    self.log_system(f"[META] CHAPTER_START: {i}")

                    # 1. Получаем словарь с текстом фактов и ID источников
                    evidence_data = self.writer._gather_evidence_for_chapter(chapter)
                    evidence_text = evidence_data["text"]
                    source_ids = evidence_data["source_ids"]
                    
                    # Инициализируем переменные перед циклом попыток
                    final_chapter_text = ""
                    feedback = None 

                    for revision_attempt in range(3):
                        if revision_attempt == 0:
                            # Первая попытка: пишем с нуля
                            current_text = self.writer.draft_section(
                                chapter, 
                                context_manager.get(), 
                                master_plan.chapters[i+1].title if i+1 < len(master_plan.chapters) else "Заключение",
                                evidence_text  
                            )
                        else:
                            # Повторные попытки: исправляем, если есть отзыв (feedback)
                            if feedback and final_chapter_text:
                                current_text = self.writer.revise_text(final_chapter_text, feedback)
                            else:
                                # Если вдруг feedback нет (защита от сбоя), используем предыдущий текст
                                current_text = final_chapter_text

                        # Проверка на явный сбой генерации
                        if current_text.strip().startswith("*["): 
                            self.log_system(f"⚠️ Сбой генерации на попытке {revision_attempt+1}, пробую еще раз...")
                            continue # Пробуем следующую итерацию
                        
                        # Полировка и соединение
                        polished_text = self.style_agent.humanize_text(current_text)
                        bridged_text = self.continuity.bridge_chapters(context_manager.get(), polished_text) if i > 0 else polished_text
                        
                        # Сохраняем текущий результат как "финальный" для этой итерации
                        final_chapter_text = bridged_text 

                        # Оценка критиком
                        feedback = self.critic.evaluate_draft(bridged_text, chapter.core_thesis)
                        
                        # Если критик доволен — выходим из цикла ревизий досрочно
                        if feedback.is_approved: 
                            self.log_system("✨ Текст принят критиком.")
                            break
                    
                    # Если после всех попыток текста нет (все крашнулись), выбрасываем ошибку
                    if not final_chapter_text:
                        raise RuntimeError(f"Writer failed to generate chapter '{chapter.title}' after 3 attempts.")

                    self.log_system(f"✅ Глава '{chapter.title}' утверждена.")

                    # 3. Создаем узел для главы и ПОЛУЧАЕМ ЕГО ID
                    chapter_node_id = self.tools.graph.add_node(
                        f"НАПИСАННЫЙ ТЕКСТ ГЛАВЫ '{chapter.title}':\n{final_chapter_text}", 
                        "written_chapter", 
                        meta={"title": chapter.title, "index": i}
                    )
                    
                    # 4. Создаем связи (ребра) от главы к ее источникам
                    if chapter_node_id and source_ids:
                        self.log_system(f"Создание связей для главы '{chapter.title}' с {len(source_ids)} источниками...")
                        for source_id in source_ids:
                            self.tools.graph.add_edge(chapter_node_id, source_id, relation="cites_source")
                    
                    # 5. Сохраняем граф с новыми узлами И СВЯЗЯМИ
                    self.tools.graph.save()
                    
                    summary = self.engine.generate_text(PromptsLibrary.CONTEXT_SUMMARY_PROTOCOL, f"ТЕКСТ:\n{final_chapter_text[:6000]}", ComplexityLevel.ROUTINE)
                    context_manager.add(chapter.title, summary)
                    
                    self.state.chapter_summaries.append(f"Глава {i+1}: {summary}")
                    
                    clean_title = re.sub(r'^\d+[\.\)]\s*', '', chapter.title).strip()
                    self.state.text_parts.append(f"## {clean_title}\n\n{final_chapter_text}\n\n")

                    if chart_path := self.writer.generate_visual_content(final_chapter_text, chapter.title):
                         self.state.text_parts.append(f"![Рисунок: {chapter.title}]({chart_path})\n\n")

                    self.state.current_chapter_index = i + 1
                    self._save_state()
                    self.tg.send_chapter_update(i + 1, len(master_plan.chapters), chapter.title)
                
                self.state.status = "writing_done"
                self._save_state()

            # ==========================================================================
            # ЭТАП 4: ФИНАЛИЗАЦИЯ И СБОРКА
            # ==========================================================================
            if self.state.status == "writing_done":
                self.log_system("\n--- ЭТАП 4: ФИНАЛЬНАЯ СБОРКА ---")
                
                full_summary_text = "\n\n".join(self.state.chapter_summaries)
                final_blocks_text = self.engine.generate_text(
                    PromptsLibrary.PUBLISHER_DIRECTIVE, 
                    f"КРАТКОЕ СОДЕРЖАНИЕ ВСЕХ НАПИСАННЫХ ГЛАВ:\n{full_summary_text}"
                )

                raw_full_text = "\n\n".join(self.state.text_parts) + "\n\n" + final_blocks_text
                
                clean_final_text = self.engine.generate_text(PromptsLibrary.FINAL_ASSEMBLER_PROTOCOL, f"ОТРЕДАКТИРУЙ:\n---\n{raw_full_text}\n---")
                final_text_for_saving = clean_final_text if len(clean_final_text) > len(raw_full_text) * 0.7 else raw_full_text

                safe_title = re.sub(r'[<>:"/\\|?*]', '', master_plan.main_title).strip()[:60]
                date_tag = datetime.now().strftime("%Y-%m-%d")
                docx_filename = f"{safe_title} ({date_tag}).docx"
                docx_filepath = os.path.join(self.config.ARTIFACTS_DIR, docx_filename)
                
                if not self.tools.doc_forge.create_styled_docx(master_plan, [final_text_for_saving], self.state.bibliography, docx_filepath):
                    raise RuntimeError("Критическая ошибка: Не удалось сохранить финальный DOCX файл.")

                # Финальный успех!
                elapsed_time = time.time() - start_time

                self.log_system(f"✅ МИССИЯ УСПЕШНО ЗАВЕРШЕНА!")
                self.log_system(f"   - Итоговый документ: {docx_filepath}")
                self.log_system(f"   - Затраченное время: {elapsed_time/60:.1f} мин.")
                
                self.tg.send_final_summary(master_plan.main_title, docx_filepath, elapsed_time)
                
                # Очищаем состояние после успешного завершения
                state_path = self._get_state_path(topic)
                if os.path.exists(state_path): os.remove(state_path)
                    
                return docx_filepath

        except Exception as e:
            # ==========================================================================
            # ОБРАБОТКА КРИТИЧЕСКИХ ОШИБОК
            # ==========================================================================
            error_message = f"❌ КРИТИЧЕСКИЙ СБОЙ МИССИИ: {e}"
            self.log_system(error_message)
            logger.error(error_message, exc_info=True) # Записываем полный traceback в файл лога
            
            # Сохраняем состояние "как есть" на момент сбоя для анализа.
            if self.state:
                self.state.status = f"error_at_{self.state.status}"
                self._save_state()
            
            self.tg.send_error(str(e), traceback.format_exc())
            return error_message

# ==============================================================================
# 22. GENESIS GUI: TACTICAL COMMAND DASHBOARD v4.0 (Fixed Progress) 
# ==============================================================================

class LogHandler(logging.Handler):
    """
    Обработчик логов, который парсит спец-теги для GUI.
    Версия 2.0: Усиленный парсинг и защита от ошибок.
    """
    def __init__(self, queue):
        super().__init__()
        self.queue = queue

    def emit(self, record):
        msg = self.format(record)
        event_type = "log"
        agent = None
        meta = {}
        
        # Определение агента для визуализации
        if "Architect" in record.name: agent = "ARCHITECT"
        elif "Researcher" in record.name: agent = "RESEARCHER"
        elif "Writer" in record.name: agent = "WRITER"
        elif "Style" in record.name: agent = "WRITER"
        elif "Critic" in record.name: agent = "CRITIC"
        elif "SYSTEM" in msg: agent = "SYSTEM"
        
        # === ПАРСИНГ ДАННЫХ ===
        
        # 1. Телеметрия (Токены и запросы)
        if "[TELEMETRY]" in msg:
            try:
                import ast
                data_str = msg.split("[TELEMETRY]")[1].strip()
                data = ast.literal_eval(data_str)
                meta["telemetry"] = data
                self.queue.put({"type": "telemetry", "meta": meta})
                return 
            except: pass

        # 2. Структура Глав
        if "[STRUCTURE]" in msg:
            try:
                data_str = msg.split("[STRUCTURE]")[1].strip()
                data = json.loads(data_str)
                meta["chapters_list"] = data
                self.queue.put({"type": "structure", "meta": meta})
                return 
            except: pass

        # 3. Текущая глава
        if "[META] CHAPTER_START:" in msg:
            try:
                idx = int(msg.split(":")[-1].strip())
                meta["current_chapter_idx"] = idx
            except: pass
            
        # 4. ИСТОЧНИКИ (ИСПРАВЛЕНО)
        # Теперь мы ловим реальные фразы, которые пишет ResearcherAgent
        if "Найдены релевантные факты" in msg or \
           "Найден и обработан источник" in msg or \
           "Прочитана страница" in msg or \
           "Добавлен узел [broad_scan_source]" in msg:
             meta["telemetry"] = {"src": 1}

        if "Готово:" in msg: event_type = "finish"

        # Отправляем сообщение
        self.queue.put({
            "type": event_type,
            "text": msg,
            "agent": agent,
            "meta": meta
        })

class HyperCortex(ctk.CTkCanvas):
    """
    3D-Визуализация мышления (Вращающийся Икосаэдр).
    """
    def __init__(self, master, width=300, height=250, **kwargs):
        super().__init__(master, width=width, height=height, bg=C_BG, highlightthickness=0, **kwargs)
        self.width = width
        self.height = height
        
        # 3D Настройки
        self.angle_x = 0
        self.angle_y = 0
        self.angle_z = 0
        self.scale = 50 # Уменьшили, чтобы влезало
        
        # Состояние
        self.running = False
        self.target_color = C_CYAN
        self.spin_speed = {"x": 0.02, "y": 0.02, "z": 0.01}
        self.pulse_speed = 0
        self.pulse_phase = 0
        
        # Геометрия Икосаэдра
        phi = (1 + math.sqrt(5)) / 2
        self.vertices = [
            [-1,  phi, 0], [ 1,  phi, 0], [-1, -phi, 0], [ 1, -phi, 0],
            [ 0, -1,  phi], [ 0,  1,  phi], [ 0, -1, -phi], [ 0,  1, -phi],
            [ phi, 0, -1], [ phi, 0,  1], [-phi, 0, -1], [-phi, 0,  1]
        ]
        self.edges = [
            [0, 11], [0, 5], [0, 1], [0, 7], [0, 10], [1, 5], [1, 9], [1, 8], [1, 7],
            [2, 11], [2, 10], [2, 6], [2, 4], [2, 3], [3, 4], [3, 9], [3, 8], [3, 6],
            [4, 5], [4, 9], [4, 11], [5, 11], [5, 9], [6, 7], [6, 8], [6, 10],
            [7, 8], [7, 10], [8, 9], [10, 11]
        ]

    def set_mode(self, agent):
        if agent == "ARCHITECT":
            self.target_color = C_GOLD
            self.spin_speed = {"x": 0.005, "y": 0.01, "z": 0.005}
            self.pulse_speed = 0.02
        elif agent == "RESEARCHER":
            self.target_color = C_BLUE
            self.spin_speed = {"x": 0, "y": 0.15, "z": 0}
            self.pulse_speed = 0
        elif agent == "WRITER":
            self.target_color = C_GREEN
            self.spin_speed = {"x": 0.03, "y": 0.03, "z": 0.03}
            self.pulse_speed = 0.15 
        elif agent == "CRITIC":
            self.target_color = C_RED
            self.spin_speed = {"x": 0.08, "y": -0.08, "z": 0.08}
            self.pulse_speed = 0.05
        else:
            self.target_color = C_CYAN
            self.spin_speed = {"x": 0.01, "y": 0.01, "z": 0.01}
            self.pulse_speed = 0.01

    def start(self):
        self.running = True
        self.animate()

    def stop(self):
        self.running = False

    def _project(self, x, y, z):
        current_scale = self.scale + math.sin(self.pulse_phase) * (self.scale * 0.1)
        factor = 300 / (300 + z) 
        px = x * current_scale * factor + self.width / 2
        py = y * current_scale * factor + self.height / 2
        return px, py

    def _rotate(self, x, y, z):
        qy = y * math.cos(self.angle_x) - z * math.sin(self.angle_x)
        qz = y * math.sin(self.angle_x) + z * math.cos(self.angle_x)
        y, z = qy, qz
        qx = x * math.cos(self.angle_y) - z * math.sin(self.angle_y)
        qz = x * math.sin(self.angle_y) + z * math.cos(self.angle_y)
        x, z = qx, qz
        qx = x * math.cos(self.angle_z) - y * math.sin(self.angle_z)
        qy = x * math.sin(self.angle_z) + y * math.cos(self.angle_z)
        x, y = qx, qy
        return x, y, z

    def animate(self):
        if not self.running: return
        self.delete("all")
        self.angle_x += self.spin_speed["x"]
        self.angle_y += self.spin_speed["y"]
        self.angle_z += self.spin_speed["z"]
        self.pulse_phase += self.pulse_speed
        
        projected = []
        for v in self.vertices:
            x, y, z = self._rotate(v[0], v[1], v[2])
            px, py = self._project(x, y, z)
            projected.append((px, py))
        
        for edge in self.edges:
            p1 = projected[edge[0]]
            p2 = projected[edge[1]]
            self.create_line(p1[0], p1[1], p2[0], p2[1], fill=self.target_color, width=1.5)
            
        for p in projected:
            self.create_oval(p[0]-2, p[1]-2, p[0]+2, p[1]+2, fill="white", outline="")

        self.after(20, self.animate)

class AgentCard(ctk.CTkFrame):
    def __init__(self, master, name, color, icon_char):
        super().__init__(master, fg_color=C_PANEL, border_width=1, border_color=C_BORDER)
        self.base_color = color
        self.indicator = ctk.CTkFrame(self, width=5, fg_color="#333", corner_radius=0)
        self.indicator.pack(side="left", fill="y")
        self.label = ctk.CTkLabel(self, text=f"[{icon_char}] {name}", font=("Orbitron", 11, "bold"), text_color="#666")
        self.label.pack(side="left", padx=10, pady=10)
        self.status_dot = ctk.CTkLabel(self, text="●", text_color="#333", font=("Arial", 16))
        self.status_dot.pack(side="right", padx=10)

    def set_active(self, active):
        if active:
            self.indicator.configure(fg_color=self.base_color)
            self.label.configure(text_color="white")
            self.status_dot.configure(text_color=self.base_color)
            self.configure(border_color=self.base_color)
        else:
            self.indicator.configure(fg_color="#333")
            self.label.configure(text_color="#666")
            self.status_dot.configure(text_color="#333")
            self.configure(border_color=C_BORDER)


# ==============================================================================
# 23. GENESIS GUI: MISSION CONTROL CENTER 
# ==============================================================================

class StatsPanel(ctk.CTkFrame):
    """Панель телеметрии."""
    def __init__(self, master, **kwargs):
        super().__init__(master, fg_color=C_PANEL, border_width=1, border_color=C_BORDER, **kwargs)
        self.grid_columnconfigure((0, 1), weight=1)
        
        self.tokens_in = self._add_stat("TOKENS IN", "0", 0, 0)
        self.tokens_out = self._add_stat("TOKENS OUT", "0", 0, 1)
        self.req_count = self._add_stat("REQUESTS", "0", 1, 0)
        self.sources = self._add_stat("SOURCES", "0", 1, 1)
        
    def _add_stat(self, title, value, row, col):
        f = ctk.CTkFrame(self, fg_color="transparent")
        f.grid(row=row, column=col, sticky="ew", padx=10, pady=5)
        ctk.CTkLabel(f, text=title, font=("Consolas", 10), text_color="#666").pack(anchor="w")
        val_lbl = ctk.CTkLabel(f, text=value, font=("Consolas", 14, "bold"), text_color=C_CYAN)
        val_lbl.pack(anchor="w")
        return val_lbl

    def update_stats(self, stats):
        if "tokens_in" in stats: self.tokens_in.configure(text=str(stats["tokens_in"]))
        if "tokens_out" in stats: self.tokens_out.configure(text=str(stats["tokens_out"]))
        if "requests" in stats: self.req_count.configure(text=str(stats["requests"]))
        if "sources" in stats: self.sources.configure(text=str(stats["sources"]))

class ChaptersList(ctk.CTkScrollableFrame):
    """Интерактивный список глав."""
    def __init__(self, master, **kwargs):
        super().__init__(master, fg_color="transparent", label_text="MISSION STRUCTURE", label_font=("Orbitron", 12), **kwargs)
        self.items = []

    def set_chapters(self, chapters):
        # Очистка
        for item in self.items: item.destroy()
        self.items = []
        
        for i, title in enumerate(chapters):
            lbl = ctk.CTkLabel(
                self, 
                text=f"{i+1}. {title}", 
                anchor="w", 
                font=("Consolas", 11),
                text_color="#888"
            )
            lbl.pack(fill="x", pady=2)
            self.items.append(lbl)

    def set_active(self, index):
        for i, item in enumerate(self.items):
            if i == index:
                item.configure(text_color=C_CYAN, font=("Consolas", 12, "bold"))
            elif i < index:
                item.configure(text_color=C_GREEN) # Пройдено
            else:
                item.configure(text_color="#888")

class GenesisDashboard(ctk.CTk):

    def __init__(self):
        super().__init__()
        self.title("GENESIS // MISSION CONTROL")
        self.geometry("1600x900")
        self.configure(fg_color=C_BG)
        
        self.queue = queue.Queue()
        self.setup_logging()
        
        # Layout
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(1, weight=1)
        
        # 1. Header
        self._build_header()
        
        # 2. Left Panel (Status & Viz)
        self.left = ctk.CTkFrame(self, width=300, fg_color="transparent")
        self.left.grid(row=1, column=0, sticky="nsew", padx=10, pady=10)
        
        self.viz = HyperCortex(self.left, width=280, height=250)
        self.viz.pack(pady=(0, 10))
        
        self.stats = StatsPanel(self.left)
        self.stats.pack(fill="x", pady=10)
        
        self.chapters_list = ChaptersList(self.left, height=300)
        self.chapters_list.pack(fill="both", expand=True)

        # 3. Center Panel (Log)
        self.center = ctk.CTkFrame(self, fg_color=C_PANEL, border_width=1, border_color=C_BORDER)
        self.center.grid(row=1, column=1, sticky="nsew", padx=10, pady=10)
        
        ctk.CTkLabel(self.center, text=" SYSTEM TERMINAL", font=("Consolas", 10), text_color="#666").pack(anchor="w", padx=5, pady=2)
        self.log_box = ctk.CTkTextbox(self.center, font=FONT_MONO, fg_color="#0F0F12", text_color=C_TEXT)
        self.log_box.pack(fill="both", expand=True, padx=2, pady=2)
        self.log_box.configure(state="disabled")

        # 4. Bottom Panel (Input)
        self.bot = ctk.CTkFrame(self, height=60, fg_color="transparent")
        self.bot.grid(row=2, column=0, columnspan=2, sticky="ew", padx=20, pady=20)
        
        self.entry = ctk.CTkEntry(self.bot, placeholder_text="Enter mission directive...", height=40, font=("Arial", 14), border_color=C_BORDER)
        self.entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.entry.bind("<Return>", lambda e: self.start())
        
        self.btn = ctk.CTkButton(self.bot, text="ENGAGE", font=("Orbitron", 11, "bold"), width=120, height=40, fg_color=C_CYAN, text_color="black", command=self.start)
        self.btn.pack(side="right")

        # Telemetry State
        self.telemetry = {"tokens_in": 0, "tokens_out": 0, "requests": 0, "sources": 0}
        
        self.after(100, self.update_gui)

    def _build_header(self):
        h = ctk.CTkFrame(self, height=50, fg_color="transparent")
        h.grid(row=0, column=0, columnspan=2, sticky="ew", padx=20, pady=(10,0))
        ctk.CTkLabel(h, text="GENESIS ", font=("Orbitron", 24, "bold"), text_color="white").pack(side="left")
        ctk.CTkLabel(h, text="Gnosis", font=("Consolas", 12), text_color=C_GREEN).pack(side="left", padx=0, pady=(10,0))

    def setup_logging(self):
        # Очистка старых хендлеров чтобы не дублировать
        logger.handlers = [h for h in logger.handlers if not isinstance(h, LogHandler)]
        h = LogHandler(self.queue)
        h.setFormatter(logging.Formatter('%(asctime)s | %(message)s', datefmt='%H:%M:%S'))
        logger.addHandler(h)

    def update_gui(self):
        try:
            while True:
                d = self.queue.get_nowait()
                msg_type = d.get("type")
                
                # === ОБРАБОТКА ТЕЛЕМЕТРИИ ===
                if msg_type == "telemetry":
                    m = d.get("meta", {}).get("telemetry", {})
                    if "req" in m: self.telemetry["requests"] += m["req"]
                    if "in" in m: self.telemetry["tokens_in"] += m["in"]
                    if "out" in m: self.telemetry["tokens_out"] += m["out"]
                    if "src" in m: self.telemetry["sources"] += m["src"]
                    self.stats.update_stats(self.telemetry)
                    continue # Не пишем в лог

                # === ОБРАБОТКА СТРУКТУРЫ ===
                if msg_type == "structure":
                    m = d.get("meta", {}).get("chapters_list", [])
                    self.chapters_list.set_chapters(m)
                    continue

                # === ОБЫЧНЫЙ ЛОГ ===
                # Логи
                if d.get("text"):
                    self.log_box.configure(state="normal")
                    self.log_box.insert("end", f"{d['text']}\n")
                    self.log_box.see("end")
                    self.log_box.configure(state="disabled")
                
                # Визуализация агентов
                if d.get("agent"):
                    self.viz.set_mode(d["agent"])
                
                # Мета-данные (структура и прогресс) из обычных логов
                m = d.get("meta", {})
                
                # Подсветка текущей главы
                if "current_chapter_idx" in m:
                    self.chapters_list.set_active(m["current_chapter_idx"])
                
                # Источники из текстовых логов (дублируем логику на всякий случай)
                if "telemetry" in m:
                    t = m["telemetry"]
                    if "src" in t: 
                        self.telemetry["sources"] += 1
                        self.stats.update_stats(self.telemetry)

                if msg_type == "finish":
                    self.btn.configure(state="normal", text="NEW MISSION")
                    self.entry.configure(state="normal")
                    self.viz.stop()

        except queue.Empty: pass
        self.after(50, self.update_gui)

    def start(self):
        t = self.entry.get()
        if not t: return
        self.entry.configure(state="disabled")
        self.btn.configure(state="disabled", text="RUNNING...")
        
        # Сброс
        self.telemetry = {"tokens_in": 0, "tokens_out": 0, "requests": 0, "sources": 0}
        self.stats.update_stats(self.telemetry)
        self.chapters_list.set_chapters([])
        
        self.viz.set_mode("SYSTEM")
        self.viz.start()
        threading.Thread(target=self._run, args=(t,), daemon=True).start()

    def _run(self, t):
        try: 
            # Создаем экземпляр и сохраняем в переменную
            system = GenesisSystem()
            system.produce_article(t)
        except Exception as e: 
            error_msg = f"CRITICAL FAILURE: {e}"
            logger.error(error_msg, exc_info=True)
            self.queue.put({"type": "finish", "text": error_msg, "agent": "SYSTEM"})
            # Убрана попытка отправки в телеграм через несуществующую переменную,
            # так как system может быть не инициализирован при ошибке создания

if __name__ == "__main__":
    app = GenesisDashboard()
    app.mainloop()
